# -*- coding: utf-8 -*-
"""
Importa lo storico dalla cartella indicata in Impostazioni (SOLO LETTURA,
mai modificata):
- clients.json           -> tabella clients
- fatture .docx (+pdf)   -> tabella invoices/items
- Fatturato 2022/23.xlsx -> tabella legacy_year_totals (totali per cliente)
"""
import os
import re
import json
import glob
import datetime

from docx import Document
import openpyxl

from . import db
from .money import parse_amount

YEAR_RE = re.compile(r'^(19|20)\d{2}$')
DATE_RE = re.compile(r'^(\d{1,2})[-./](\d{1,2})[-./](\d{2,4})$')
# I numeri fattura reali finora sono < 200; oltre 400 e' quasi certamente un errore di nome file
MAX_SANE_NUMBER = 400


def trailing_number(name):
    # gruppi di 2-3 cifre NON parte di numeri piu' lunghi (esclude anni tipo '2022')
    nums = re.findall(r'(?<!\d)(\d{2,3})(?!\d)', os.path.splitext(os.path.basename(name))[0])
    n = int(nums[-1]) if nums else None
    if n is not None and n > MAX_SANE_NUMBER:
        return None
    return n


def client_from_filename(name):
    stem = os.path.splitext(os.path.basename(name))[0]
    return re.split(r'[#^_]| \d', stem)[0].strip(' -_^N') or stem


def parse_date_iso(t):
    """'02-06-26' / '16.01.24' / '02/06/2026' -> '2026-06-02' (None se non valida)."""
    m = DATE_RE.match(t.strip())
    if not m:
        return None
    dd, mm, yy = int(m.group(1)), int(m.group(2)), int(m.group(3))
    if yy < 100:
        yy += 2000
    try:
        return datetime.date(yy, mm, dd).isoformat()
    except ValueError:
        return None


def find_year_dirs(root):
    return sorted(d for d in os.listdir(root)
                  if YEAR_RE.match(d) and os.path.isdir(os.path.join(root, d)))


def extract_docx(path):
    """Estrae (numero, cliente, indirizzo, data_iso, items, total_cents) da una fattura .docx.
    items = [(qty, desc, unit_cents, total_cents), ...]"""
    d = Document(path)
    number = trailing_number(path)
    client = client_from_filename(path)
    address = ''
    date_iso = None
    items = []
    total_cents = None
    try:
        header = d.tables[0]
        for p in header.rows[0].cells[1].paragraphs:
            t = p.text.strip()
            iso = parse_date_iso(t)
            if iso:
                date_iso = iso
        lines = [x.strip() for x in header.rows[1].cells[0].text.split('\n') if x.strip()]
        if lines:
            client = lines[0]
            address = '\n'.join(lines[1:])
    except Exception:
        pass
    # righe articolo: tabella 2, righe dati 1..n (qty | descrizione | unit | totale)
    try:
        t2 = d.tables[2]
        for row in t2.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 4:
                continue
            qty, desc, unit, tot = cells[0], cells[1], cells[2], cells[3]
            if not desc and not tot:
                continue
            items.append((qty or '1', desc,
                          parse_amount(unit), parse_amount(tot)))
    except Exception:
        pass
    # totale: tabella 3 (TOTAL DUE), ultima cella
    try:
        total_cents = parse_amount(d.tables[3].rows[0].cells[-1].text)
    except Exception:
        pass
    if total_cents is None and items:
        vals = [t for (_, _, _, t) in items if t is not None]
        total_cents = sum(vals) if vals else None
    return number, client, address, date_iso, items, total_cents


def _dedupe_key(path):
    return (trailing_number(path), client_from_filename(path).lower()[:6])


def import_all(con, root, log=None):
    """Import completo. Idempotente: cancella e ricrea i record con source='import'."""
    msgs = []

    def say(m):
        msgs.append(m)
        if log:
            log(m)

    # --- clienti dal registro ---
    reg_path = os.path.join(root, 'clients.json')
    if os.path.exists(reg_path):
        reg = json.load(open(reg_path, encoding='utf-8'))
        n = 0
        for key, c in reg.items():
            if key.startswith('_'):
                continue
            addr = c.get('address', [])
            con.execute(
                'INSERT INTO clients(key,name,address1,address2,file_label,notes) '
                'VALUES(?,?,?,?,?,?) ON CONFLICT(key) DO UPDATE SET '
                'name=excluded.name, address1=excluded.address1, '
                'address2=excluded.address2, file_label=excluded.file_label',
                (key, c.get('name', key),
                 addr[0] if len(addr) > 0 else '',
                 '\n'.join(addr[1:]) if len(addr) > 1 else '',
                 c.get('file_label', ''), ''))
            n += 1
        say(f'Clienti importati dal registro: {n}')

    # --- pulizia import precedente (idempotenza) ---
    con.execute("DELETE FROM items WHERE invoice_id IN "
                "(SELECT id FROM invoices WHERE source='import')")
    con.execute("DELETE FROM invoices WHERE source='import'")
    con.execute('DELETE FROM legacy_year_totals')

    clients = {r['id']: r for r in con.execute('SELECT * FROM clients')}

    def match_client_id(name):
        low = (name or '').lower()
        for cid, c in clients.items():
            cn = c['name'].lower()
            if cn and (low.startswith(cn) or cn.startswith(low.split()[0] if low.split() else low)):
                return cid
        return None

    this_year = datetime.date.today().year
    seen = set()
    count = 0
    for y in find_year_dirs(root):
        ydir = os.path.join(root, y)
        for path in sorted(glob.glob(os.path.join(ydir, '**', '*.docx'), recursive=True)):
            base = os.path.basename(path)
            if base.startswith(('~$', '.')):
                continue
            try:
                number, client, address, date_iso, items, total = extract_docx(path)
            except Exception as e:
                say(f'  ! Illeggibile: {y}/{base} ({e})')
                continue
            if number is not None and number in (int(y), int(y) % 100):
                number = None  # il "numero" era in realta' l'anno nel nome file
            key = _dedupe_key(path)
            if key in seen:
                continue
            seen.add(key)
            status = 'pagata' if int(y) < this_year else 'emessa'
            cur = con.execute(
                'INSERT INTO invoices(number, client_id, client_name, client_address, date, '
                'year, total_cents, status, source, source_file, created_at) '
                "VALUES(?,?,?,?,?,?,?,?, 'import', ?, ?)",
                (number, match_client_id(client), client, address,
                 date_iso, int(y), total, status,
                 os.path.relpath(path, root), db.now_iso()))
            inv_id = cur.lastrowid
            for pos, (qty, desc, unit_c, tot_c) in enumerate(items):
                con.execute('INSERT INTO items(invoice_id,pos,qty,description,unit_cents,total_cents) '
                            'VALUES(?,?,?,?,?,?)', (inv_id, pos, qty, desc, unit_c, tot_c))
            count += 1
        # PDF senza docx corrispondente (solo dal 2024: per il 2022/23 fanno fede gli Excel)
        if int(y) < 2024:
            continue
        for path in sorted(glob.glob(os.path.join(ydir, '**', '*.pdf'), recursive=True)):
            base = os.path.basename(path)
            if base.startswith(('~$', '.')):
                continue
            key = _dedupe_key(path)
            if key in seen:
                continue
            seen.add(key)
            number = trailing_number(path)
            if number is not None and number in (int(y), int(y) % 100):
                number = None
            client = client_from_filename(path)
            status = 'pagata' if int(y) < this_year else 'emessa'
            con.execute(
                'INSERT INTO invoices(number, client_id, client_name, date, year, total_cents, '
                "status, source, source_file, notes, created_at) "
                "VALUES(?,?,?,?,?,?,?, 'import', ?, 'Solo PDF: importo da verificare', ?)",
                (number, match_client_id(client), client, None, int(y), None,
                 status, os.path.relpath(path, root), db.now_iso()))
            count += 1
    say(f'Fatture importate dalle cartelle anno: {count}')

    # --- Excel storici 2022/2023: totali per cliente (fonte autorevole per quegli anni) ---
    for y in (2022, 2023):
        xp = os.path.join(root, str(y), f'Fatturato {y}.xlsx')
        if not os.path.exists(xp):
            continue
        from .db import get_settings
        rows = _parse_legacy_xlsx(xp, get_settings(con).get('business_name', ''))
        for client, invoiced, paid, note in rows:
            con.execute('INSERT INTO legacy_year_totals(year, client, invoiced_cents, paid_cents, note) '
                        'VALUES(?,?,?,?,?)', (y, client, invoiced, paid, note))
        tot = sum(r[1] or 0 for r in rows)
        say(f'Storico {y} da Excel: {len(rows)} clienti, totale fatturato {tot/100:.2f} CHF')

    # --- ri-applica le correzioni manuali: il reimport non le deve perdere ---
    from .corrections import apply_all
    n_corr = apply_all(con)
    if n_corr:
        say(f'Correzioni manuali ri-applicate: {n_corr}')

    con.commit()
    return msgs


def _parse_legacy_xlsx(path, mio_nome=''):
    """Legge i 'Fatturato YYYY.xlsx': blocchi per cliente con Total Invoiced / Amount Paid.

    `mio_nome` e' il nome dell'attivita': in cima a quei fogli c'e' spesso
    l'intestazione di chi li ha fatti, e non e' un cliente."""
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.worksheets[0]
    rows = list(ws.iter_rows(values_only=True))
    out = []
    current = None
    invoiced = paid = None
    note = ''
    skip_names = {'invoice details', 'summary', 'total invoiced:', 'total paid:'}
    mio_nome = (mio_nome or '').strip().lower()

    def flush():
        nonlocal current, invoiced, paid, note
        if current and invoiced is not None:
            out.append((current, _cents(invoiced), _cents(paid), note))
        current, invoiced, paid, note = None, None, None, ''

    for r in rows:
        a = r[0]
        b = r[1] if len(r) > 1 else None
        if isinstance(a, str):
            t = a.strip()
            low = t.lower()
            if low == 'total invoiced' and isinstance(b, (int, float)):
                invoiced = b
            elif low == 'amount paid' and isinstance(b, (int, float)):
                paid = b
            elif (low not in skip_names and not low.startswith(('total', 'amount', '#', '('))
                  and not (mio_nome and low.startswith(mio_nome))
                  and isinstance(t, str) and len(t) > 2 and not t[0].isdigit()):
                flush()
                current = t
        elif a is None and current is None:
            continue
    flush()
    return out


def _cents(v):
    if v is None:
        return None
    from decimal import Decimal, ROUND_HALF_UP
    return int((Decimal(str(v)) * 100).quantize(Decimal('1'), rounding=ROUND_HALF_UP))
