# -*- coding: utf-8 -*-
"""Genera la fattura .docx riempiendo il modello in assets/fattura-modello.docx."""
import os
import re
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from .money import fmt_dash
from . import language as L
from . import branding


# Le condizioni di pagamento le scrive chi usa l'app, nelle Impostazioni. Se
# non le ha toccate, quella che si vede e' la frase di partenza dell'app, e
# quella si puo' tradurre; se invece l'ha riscritta, sono parole sue e si
# lasciano stare — tradurgliele sarebbe metterle in bocca qualcosa che non ha
# detto.
CONDIZIONI_DI_SERIE = 'Payable within 30 days net to:'


def condizioni(settings, lingua=None):
    testo = (settings or {}).get('terms', '')
    if testo.strip() == CONDIZIONI_DI_SERIE:
        return L.t_doc('Pagabile entro 30 giorni netti a:', lingua)
    return testo

APP_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE = os.path.join(APP_DIR, 'assets', 'fattura-modello.docx')


def _set_paragraph_text(p, text):
    """Riscrive un paragrafo, spazzando via tutto quello che c'era.

    Word a volte tiene pezzi di testo dentro «controlli contenuto» (w:sdt),
    che python-docx non conta fra i run del paragrafo. Fermarsi ai run lascia
    quei pezzi nel documento: non si vedono aprendo il file, ma nel .docx ci
    sono. E' cosi' che il nome di chi aveva fatto il template restava dentro
    ogni fattura.

    Non basta svuotarli: quello del ringraziamento e' AGGANCIATO alla proprieta'
    «Azienda» del documento, e Word lo riempirebbe di nuovo da solo alla prima
    apertura, raddoppiando il nome. Si tolgono di mezzo del tutto."""
    for sdt in list(p._p.iter(qn('w:sdt'))):
        genitore = sdt.getparent()
        if genitore is not None:
            genitore.remove(sdt)
    for t in p._p.findall('.//' + qn('w:t')):
        t.text = ''
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(text)


def _set_cell(cell, text):
    _set_paragraph_text(cell.paragraphs[0], text)


# Dove sta, dentro il template, ciascun dato di chi emette la fattura.
# La cella in alto a sinistra tiene l'intestazione, le ultime righe del corpo
# tengono condizioni di pagamento e IBAN. Nel template ci sono solo segnaposto:
# i valori veri arrivano sempre dalle Impostazioni, cosi' chi usa l'app manda
# fatture col PROPRIO nome e il PROPRIO IBAN senza dover aprire Word.
RIGHE_MITTENTE = {0: 'business_name', 1: 'business_uid', 5: 'business_addr1',
                  6: 'business_addr2', 7: 'business_phone', 8: 'business_web'}
RIGA_RINGRAZIAMENTO = 6
RIGA_CONDIZIONI = 11      # la 10 e' l'etichetta «Terms:», non il testo
RIGA_NOME_PAGAMENTO = 12
RIGA_IBAN = 13


def _scrivi_intestazione(d, settings, lingua=None):
    """Chi emette la fattura: nome, UID, indirizzo, contatti, condizioni, IBAN.

    Se qualcuno rimaneggia il template e una riga non c'e' piu', quella riga si
    salta invece di far fallire la fattura: meglio un documento con un campo in
    meno che nessun documento.
    """
    if settings is None:
        return
    cella = d.tables[0].rows[0].cells[0]
    for i, chiave in RIGHE_MITTENTE.items():
        if i < len(cella.paragraphs):
            _set_paragraph_text(cella.paragraphs[i], settings.get(chiave, ''))
    nome = settings.get('business_name', '')
    testi = {
        RIGA_RINGRAZIAMENTO: (L.t_doc('Grazie per aver scelto {nome}!', lingua)
                              .format(nome=nome) if nome else ''),
        RIGA_CONDIZIONI: condizioni(settings, lingua),
        RIGA_NOME_PAGAMENTO: nome,
        RIGA_IBAN: 'IBAN: ' + settings.get('business_iban', ''),
    }
    for i, testo in testi.items():
        if i < len(d.paragraphs):
            _set_paragraph_text(d.paragraphs[i], testo)


def _scrivi_proprieta(d, settings):
    """Le proprieta' del documento: autore, ultimo che l'ha modificato, azienda.

    Word le riempie con chi ha creato il file e se le porta dietro per sempre.
    Un cliente che apre la fattura le vede in Informazioni documento, e chi usa
    l'app spedirebbe fatture firmate da un altro senza saperlo."""
    nome = (settings or {}).get('business_name', '')
    cp = d.core_properties
    cp.author = nome
    cp.last_modified_by = nome
    for campo in ('title', 'subject', 'comments', 'category', 'keywords',
                  'content_status', 'identifier', 'version'):
        try:
            setattr(cp, campo, '')
        except (AttributeError, ValueError):     # pragma: no cover
            pass
    # <Company> sta in docProps/app.xml, che python-docx non espone
    for parte in d.part.package.iter_parts():
        if str(parte.partname) != '/docProps/app.xml':
            continue
        xml = parte.blob.decode('utf-8', 'replace')
        xml = re.sub(r'<Company>.*?</Company>', f'<Company>{nome}</Company>', xml)
        xml = re.sub(r'<Manager>.*?</Manager>', '<Manager></Manager>', xml)
        parte._blob = xml.encode('utf-8')


def _scrivi_logo(d):
    """Sostituisce il logo del template con quello dell'utente.

    Lo spazio del logo nel documento Word ha misure fisse: un logo di forma
    diversa ci finirebbe stirato. Per questo lo consegniamo gia' centrato
    dentro una cornice trasparente della forma giusta."""
    for rel in d.part.rels.values():
        if 'image' not in rel.reltype:
            continue
        parte = rel.target_part
        try:
            larghezza, altezza = parte.image.px_width, parte.image.px_height
        except Exception:
            larghezza, altezza = 388, 288
        parte._blob = branding.adattato(larghezza, altezza)
        return


# Le etichette che stanno gia' scritte dentro il modello Word. Il modello e'
# in inglese, che e' la lingua con cui l'app e' nata; per un cliente tedesco
# o italiano vanno riscritte una per una, perche' Word non ha segnaposti.
ETICHETTE = {
    'QUANTITY': 'QUANTITÀ',
    'DESCRIPTION': 'DESCRIZIONE',
    'UNIT PRICE': 'PREZZO UNITARIO',
    'TOTAL': 'TOTALE',
    'TOTAL due': 'TOTALE DA PAGARE',
    'Terms': 'Condizioni',
}


def _traduci_etichette(d, lingua):
    """Riscrive le etichette del modello nella lingua del cliente.

    Si lavora sui pezzi di testo del file Word, non sui paragrafi: nel
    modello le intestazioni della tabella stanno dentro riquadri di testo,
    che python-docx non mostra fra i paragrafi. Passando dai pezzi si
    raggiungono tutte, e per giunta si tocca solo la parola: «Terms» resta
    sottolineato e i due punti restano dove sono.

    Va chiamata a modello appena aperto, prima che ci finiscano dentro i
    dati: cosi' non puo' scambiare per un'etichetta il nome di un cliente.
    """
    if L.normalizza_doc(lingua) == L.PREDEFINITA_DOC:
        return                       # il modello e' gia' cosi'
    for pezzo in d.element.body.iter(qn('w:t')):
        chiave = ETICHETTE.get((pezzo.text or '').strip())
        if chiave:
            pezzo.text = L.t_doc(chiave, lingua)


def build_docx(out_path, number, date_str, client_name, addr_lines, items, total_cents,
               settings=None, lingua=None):
    """items: lista di dict {qty, description, unit_cents, total_cents} (max 8 righe)."""
    d = Document(TEMPLATE)
    _traduci_etichette(d, lingua)
    _scrivi_intestazione(d, settings, lingua)
    _scrivi_logo(d)
    _scrivi_proprieta(d, settings)
    header = d.tables[0]
    hcell = header.rows[0].cells[1]
    for p in hcell.paragraphs:
        t = p.text.strip()
        if '#' in t:
            _set_paragraph_text(p, f' #{number}')
        elif re.match(r'^\d{1,2}[-./]\d{1,2}[-./]\d', t):
            _set_paragraph_text(p, f' {date_str}')
    # destinatario
    rcell = header.rows[1].cells[0]
    nonempty = [p for p in rcell.paragraphs if p.text.strip()]
    lines = [client_name] + [l for l in addr_lines if l and l.strip()]
    if nonempty:
        _set_paragraph_text(nonempty[0], lines[0])
        for p in nonempty[1:]:
            _set_paragraph_text(p, '')
        anchor = nonempty[0]._p
        for line in lines[1:]:
            new_p = deepcopy(anchor)
            anchor.addnext(new_p)
            anchor = new_p
            _set_paragraph_text(Paragraph(new_p, rcell.paragraphs[0]._parent), line)
    # righe articolo
    table = d.tables[2]
    for i, it in enumerate(items[:8]):
        row = table.rows[1 + i]
        _set_cell(row.cells[0], str(it['qty']))
        _set_cell(row.cells[1], it['description'])
        _set_cell(row.cells[2], fmt_dash(it['unit_cents']) if it['unit_cents'] is not None else '')
        _set_cell(row.cells[3], (fmt_dash(it['total_cents']) + ' CHF') if it['total_cents'] is not None else '')
    # totale
    tot = d.tables[3]
    _set_cell(tot.rows[0].cells[-1], fmt_dash(total_cents) + ' CHF')
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    d.save(out_path)
    return out_path
