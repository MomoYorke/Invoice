# -*- coding: utf-8 -*-
"""
Batteria di test sui calcoli. Nessuna dipendenza esterna: verifica il motore
matematico (centesimi interi) su casi normali e casi-trappola.
run_all() ritorna (tutti_ok, [risultati]) dove ogni risultato è
(categoria, descrizione, ok, dettaglio).
"""
import io
import os
import re
from decimal import Decimal

from .money import parse_amount, fmt_chf, fmt_dash, parse_qty, line_total
from . import verify
from . import menu


def _check(results, cat, desc, got, expected):
    ok = got == expected
    detail = '' if ok else f'ottenuto {got!r}, atteso {expected!r}'
    results.append((cat, desc, ok, detail))


def run_all():
    r = []

    # --- lettura importi in tutti i formati che potresti scrivere ---
    parse_cases = [
        ('110.-', 11000), ('110.–', 11000), ('150,00 CHF', 15000),
        ("1'800.00", 180000), ('1.800,00CHF', 180000), ('1800', 180000),
        ('1800.5', 180050), ('CHF 110', 11000), ('110 CHF', 11000),
        ('0.-', 0), ("1'800.-", 180000), ('1.800', 180000), ('150.-', 15000),
        ('24234.3', 2423430), ('1,800.00', 180000), ('2.796,8', 279680),
        ('99.95', 9995), ('  ', None), ('abc', None), (None, None),
        ('12.345,67', 1234567), ("12'345.67", 1234567), ('0', 0),
        ('0.01', 1), ('0,05', 5), ('1000000', 100000000),
    ]
    for txt, exp in parse_cases:
        _check(r, 'Lettura importi', f'"{txt}" → {fmt_chf(exp) if exp is not None else "vuoto"}',
               parse_amount(txt), exp)

    # --- moltiplicazione quantità × prezzo (il cuore della fattura) ---
    mult_cases = [
        (12, 15000, 180000), ('12', 15000, 180000), (1, 11000, 11000),
        (3, 3333, 9999), (10, 9995, 99950), (24, 15000, 360000),
        (0, 15000, 0), (100, 100, 10000),
    ]
    for qty, unit, exp in mult_cases:
        _check(r, 'Quantità × prezzo', f'{qty} × {fmt_dash(unit)} = {fmt_dash(exp)}',
               line_total(qty, unit), exp)

    # --- quantità decimali con arrotondamento corretto ---
    _check(r, 'Arrotondamento', "1.5 × 110.- = 165.-", line_total('1.5', 11000), 16500)
    _check(r, 'Arrotondamento', "0.333 × 100.- = 33.30 (half-up)",
           line_total(parse_qty('0.333'), 10000), 3330)
    _check(r, 'Arrotondamento', "2.5 × 99.99 = 249.98 (half-up sul centesimo)",
           line_total(parse_qty('2.5'), 9999), 24998)

    # --- formattazione (quello che finisce sulla fattura) ---
    fmt_cases = [
        (180000, "1'800.00 CHF"), (2423430, "24'234.30 CHF"),
        (11000, "110.00 CHF"), (0, "0.00 CHF"), (100000000, "1'000'000.00 CHF"),
    ]
    for cents, exp in fmt_cases:
        _check(r, 'Formato CHF', f'{cents} centesimi → {exp}', fmt_chf(cents), exp)

    dash_cases = [(180000, "1'800.-"), (180050, "1'800.50"), (11000, "110.-"),
                  (9995, "99.95"), (0, "0.-")]
    for cents, exp in dash_cases:
        _check(r, 'Formato fattura', f'{cents} centesimi → {exp}', fmt_dash(cents), exp)

    # --- round-trip: leggi ciò che hai scritto e riottieni lo stesso numero ---
    for cents in [11000, 15000, 180000, 2423430, 9995, 1, 100000000]:
        back = parse_amount(fmt_chf(cents))
        _check(r, 'Andata e ritorno', f'{fmt_chf(cents)} riletto = stesso valore', back, cents)

    # --- somma di più righe (nessuna perdita di centesimi) ---
    rows_totals = [180000, 0, 11000, 33330, 9995]
    _check(r, 'Somma righe', "1800 + 0 + 110 + 333.30 + 99.95 = 2343.25",
           sum(rows_totals), 234325)

    # --- coerenza: qty×unit deve dare il totale riga atteso, sempre ---
    consistency_ok = all(
        line_total(q, u) == Decimal(q) * Decimal(u) if isinstance(q, int) else True
        for q, u in [(12, 15000), (3, 3333), (24, 15000), (100, 100)])
    _check(r, 'Coerenza interna', 'quantità intere: nessun arrotondamento spurio',
           consistency_ok, True)

    _test_email(r)
    _test_oggetto(r)
    _test_intestazione_fattura(r)
    _test_marchio(r)
    _test_clienti_crediti(r)
    _test_servizi(r)
    _test_servizi_riconosciuti(r)
    _test_da_fare(r)
    _test_lingua(r)
    _test_primi_passi(r)
    _test_icone(r)
    _test_menu(r)
    _test_finestra_stretta(r)
    _test_calendario(r)

    all_ok = all(x[2] for x in r)
    return all_ok, r


class _Finta(dict):
    """Una riga di database finta: si comporta come sqlite3.Row."""
    def keys(self):
        return list(super().keys())


def _test_email(r):
    """Il testo dell'email si costruisce senza toccare la rete: si puo' provare."""
    from . import mailer
    from .db import DEFAULT_SETTINGS

    # Chi usa l'app ha scritto i suoi servizi in Impostazioni: e' il caso
    # normale. Le regole non sono piu' nel programma, quindi vanno date.
    S = dict(DEFAULT_SETTINGS,
             servizi_abbonamento='Running Coaching = running coaching\n'
                                 'Online Coaching = coaching online',
             servizi_pacchetto='Personal Training = session, personal training')

    inv = _Finta(number=99, total_cents=110000, pdf_path='', source_file='',
                 client_name='Chiara De Santis')
    mensile = _Finta(name='Chiara De Santis', email='b@esempio.ch',
                     abbonamento=1, tono='informale')
    pacchetto = _Finta(name='Petra Müller', email='d@esempio.ch',
                       abbonamento=0, tono='formale')

    m = mailer.componi(inv, mensile, S, ['Monthly abo: running coaching'])
    _check(r, 'Email', 'abbonamento: "this month\'s invoice"',
           "this month's invoice for Running Coaching" in m['body'], True)
    _check(r, 'Email', 'abbonamento: c\'è la frase sull\'ordine permanente',
           'standing order' in m['body'], True)
    _check(r, 'Email', 'tono informale: si chiude col saluto informale',
           m['body'].rstrip().endswith('Best,'), True)

    m2 = mailer.componi(inv, pacchetto, S, ['10 Sessions Pack – Personal Training'])
    _check(r, 'Email', 'pacchetto: NON dice "this month\'s"',
           "this month's" in m2['body'], False)
    _check(r, 'Email', "pacchetto: l'apertura nomina il servizio",
           'your invoice for Personal Training' in m2['body'], True)
    _check(r, 'Email', 'tono formale: si firma per esteso',
           'Best regards,' in m2['body'], True)
    _check(r, 'Email', 'niente frase sull\'ordine permanente se non è abbonato',
           'standing order' in m2['body'], False)

    _check(r, 'Email', 'nome di battesimo dal nome completo',
           mailer.nome_di_battesimo('Chiara De Santis'), 'Chiara')

    # i due modelli: sceglierli da soli e poterli forzare a mano
    prove = dict(S, email_corpo_coaching='TESTO-COACHING', email_corpo_pt='TESTO-PT')
    _check(r, 'Email', 'running coaching → modello coaching',
           mailer.modello_di(['Monthly abo: running coaching'], S), 'coaching')
    _check(r, 'Email', 'coaching online → modello coaching',
           mailer.modello_di(['Coaching online - August'], S), 'coaching')
    _check(r, 'Email', 'righe di un pacchetto → modello «pacchetto»',
           mailer.modello_di(['10 Sessions Pack – Personal Training'], S), 'pt')

    # chi fa un altro mestiere: nessuna regola riconosce le sue righe, e
    # l'email non deve inventargli un servizio che non vende
    _check(r, 'Email', 'servizio non riconosciuto: non se ne inventa uno',
           mailer.servizio_di(['Pacchetto 10 sedute di fisioterapia'], S), '')
    _check(r, 'Email', "servizio ignoto: l'apertura non lo nomina",
           mailer.componi(inv, pacchetto, S, ['Pacchetto 10 sedute'])['body'].split('\n')[2],
           'Please find attached your invoice.')
    _check(r, 'Email', "servizio noto: l'apertura lo nomina come prima",
           mailer.componi(inv, pacchetto, S, ['10 Sessions Pack'])['body'].split('\n')[2],
           'Please find attached your invoice for Personal Training.')
    _check(r, 'Email', 'il modello dedotto finisce nel testo',
           'TESTO-COACHING' in mailer.componi(inv, mensile, prove,
                                              ['Monthly abo: running coaching'])['body'], True)
    _check(r, 'Email', 'modello forzato a mano: vince sulla deduzione',
           'TESTO-PT' in mailer.componi(inv, mensile, prove,
                                        ['Monthly abo: running coaching'],
                                        modello='pt')['body'], True)
    _check(r, 'Email', 'modello inventato: si torna a quello dedotto',
           mailer.componi(inv, mensile, prove, ['Monthly abo: running coaching'],
                          modello='inesistente')['modello'], 'coaching')
    _check(r, 'Email', 'un testo scritto a mano batte tutti e due i modelli',
           'A MANO' in mailer.componi(inv, mensile, prove,
                                      ['Monthly abo: running coaching'],
                                      corpo='A MANO')['body'], True)
    _check(r, 'Email', 'senza indirizzo: lo dice invece di mandare a vuoto',
           bool(mailer.componi(inv, _Finta(name='X', email='', abbonamento=0, tono='informale'),
                               S)['problemi']), True)
    _check(r, 'Email', 'indirizzo malscritto: lo dice',
           bool(mailer.componi(inv, _Finta(name='X', email='pippo.esempio', abbonamento=0,
                                           tono='informale'), S)['problemi']), True)
    esito = mailer.spedisci({'to': 'x@y.ch', 'subject': '', 'body': '', 'allegati': []},
                            dict(S, smtp_pass=''))
    _check(r, 'Email', 'senza password non prova nemmeno a collegarsi', esito[0], False)
    _check(r, 'Email', 'il motivo del fallimento è etichettato', esito[2], 'config')
    import smtplib
    m = mailer.costruisci_messaggio({'to': 'cliente@esempio.ch', 'subject': 'x',
                                     'body': 'y', 'allegati': []},
                                    dict(S, smtp_user='io@esempio.ch'), None, copia_a_me=True)

    class _FintoSMTP(smtplib.SMTP):
        def __init__(self):
            self.inviati = []; self.esmtp_features = {}; self.does_esmtp = 0
        def ehlo_or_helo_if_needed(self):
            pass
        def sendmail(self, da, a, testo, *args, **kw):
            self.inviati.append((a, testo.decode() if isinstance(testo, bytes) else testo))
            return {}

    f = _FintoSMTP()
    smtplib.SMTP.send_message(f, m)
    a_chi, spedito = f.inviati[0]
    _check(r, 'Email', 'la copia nascosta arriva anche a te',
           'io@esempio.ch' in a_chi, True)
    _check(r, 'Email', 'il cliente riceve comunque la sua',
           'cliente@esempio.ch' in a_chi, True)
    _check(r, 'Email', 'il cliente NON vede che ti sei mandato una copia',
           'Bcc:' in spedito, False)

    _check(r, 'Email', 'la password non finisce nei messaggi d\'errore',
           'segreta' in mailer._nascondi('errore con segreta dentro',
                                         {'smtp_pass': 'segreta'}), False)


CAL_PROVA = """BEGIN:VCALENDAR
X-WR-CALNAME:Allenamenti
BEGIN:VEVENT
UID:tizio@g
DTSTART;TZID=Europe/Zurich:20260818T073000
SUMMARY:Marco
RRULE:FREQ=WEEKLY;BYDAY=TU
EXDATE;TZID=Europe/Zurich:20260901T073000
END:VEVENT
BEGIN:VEVENT
UID:caio@g
DTSTART;TZID=Europe/Zurich:20260820T093000
SUMMARY:Anna
RRULE:FREQ=WEEKLY;BYDAY=TH
END:VEVENT
BEGIN:VEVENT
UID:caio@g
RECURRENCE-ID;TZID=Europe/Zurich:20260827T093000
DTSTART;TZID=Europe/Zurich:20260827T093000
SUMMARY:Anna
STATUS:CANCELLED
END:VEVENT
BEGIN:VEVENT
UID:sempronio@g
DTSTART;TZID=Europe/Zurich:20260821T074500
SUMMARY:Sara\\, spostata
RECURRENCE-ID;TZID=Europe/Zurich:20260821T074500
END:VEVENT
END:VCALENDAR"""


def _test_oggetto(r):
    """L'oggetto cambia col servizio, e il mese si deduce invece di inventarlo."""
    from . import mailer
    from .db import DEFAULT_SETTINGS as S

    inv = _Finta(number=87, total_cents=11000, pdf_path='', source_file='',
                 client_name='Chiara De Santis')
    cli = _Finta(name='Chiara De Santis', email='b@esempio.ch',
                 abbonamento=1, tono='informale')

    # --- il periodo scritto sulla fattura ---
    _check(r, 'Oggetto email', 'periodo a cavallo di due mesi → i due mesi',
           mailer.mesi_da_descrizioni(
               ['Monthly abo: running coaching 13.08.26 \u2013 12.09.26']), (8, 9))
    _check(r, 'Oggetto email', 'periodo dentro un mese solo → un mese',
           mailer.mesi_da_descrizioni(
               ['Monthly abo: running coaching 01.08.26 \u2013 31.08.26']), (8,))
    _check(r, 'Oggetto email', 'mese scritto a parole',
           mailer.mesi_da_descrizioni(['Monthly abo: running coaching (August)']), (8,))
    _check(r, 'Oggetto email', 'niente periodo → niente mese',
           mailer.mesi_da_descrizioni(['10 Sessions Pack \u2013 Personal Training at Home']), ())
    _check(r, 'Oggetto email', 'una parola che comincia come un mese non è un mese',
           mailer.mesi_da_descrizioni(['Marathon plan, Decathlon, Augmented']), ())
    _check(r, 'Oggetto email', 'il trattino corto vale come quello lungo',
           mailer.mesi_da_descrizioni(['abo 13.08.26 - 12.09.26']), (8, 9))

    # --- come li abbrevia lui ---
    stile = mailer.stile_mesi(['Online Running Coaching \u2013 [Aug/Sept] \u2013 EM'])
    _check(r, 'Oggetto email', 'impara «Sept» da come l\'ha scritto lui', stile.get(9), 'Sept')
    _check(r, 'Oggetto email', 'senza esempi usa l\'abbreviazione di tre lettere',
           mailer.etichetta_mesi((9, 10)), 'Sep/Oct')
    _check(r, 'Oggetto email', 'con l\'esempio scrive come lui',
           mailer.etichetta_mesi((9, 10), stile), 'Sept/Oct')
    _check(r, 'Oggetto email', 'vince la forma usata per ultima',
           mailer.stile_mesi(['[Sept]', '[Sep]']).get(9), 'Sep')

    # --- il mese completo, con e senza aiuto dal passato ---
    _check(r, 'Oggetto email', 'Chiara: 13.08–12.09 → Aug/Sept',
           mailer.mese_oggetto(['Monthly abo: running coaching 13.08.26 \u2013 12.09.26'],
                               ['Online Running Coaching \u2013 [Aug/Sept] \u2013 EM']),
           'Aug/Sept')
    _check(r, 'Oggetto email', 'senza periodo: dall\'ultima mail, avanti di un mese',
           mailer.mese_oggetto(['Monthly abo: running coaching'],
                               ['Online Running Coaching \u2013 [Aug/Sept] \u2013 EM']),
           'Sept/Oct')
    _check(r, 'Oggetto email', 'dicembre passa a gennaio, non al mese 13',
           mailer.mese_oggetto(['abo'], ['[Dec]']), 'Jan')
    _check(r, 'Oggetto email', 'niente da cui dedurre → resta vuoto',
           mailer.mese_oggetto(['10 Sessions Pack'], []), '')

    # --- l'oggetto vero e proprio ---
    coaching = mailer.componi(inv, cli, S, ['abo 13.08.26 \u2013 12.09.26'],
                              modello='coaching', mese='Aug/Sept')
    _check(r, 'Oggetto email', 'abbonamento: oggetto col mese dentro',
           coaching['subject'],
           'Invoice \u2013 [Aug/Sept]')
    pt = mailer.componi(inv, cli, S, ['10 Sessions Pack'], modello='pt')
    _check(r, 'Oggetto email', 'pacchetto di sedute: oggetto senza mese',
           pt['subject'],
           'Invoice')
    _check(r, 'Oggetto email', 'cambiare servizio cambia l\'oggetto',
           coaching['subject'] != pt['subject'], True)
    senza = mailer.componi(inv, cli, S, ['abo'], modello='coaching', mese='')
    _check(r, 'Oggetto email', 'mese ignoto: resta «[month]» da riempire a mano',
           '[month]' in senza['subject'], True)
    _check(r, 'Oggetto email', 'e l\'app lo segnala', senza['mese_mancante'], True)
    _check(r, 'Oggetto email', 'col pacchetto il mese non manca mai',
           pt['mese_mancante'], False)


def _test_intestazione_fattura(r):
    """Chi emette la fattura lo dicono le Impostazioni, non il template Word.

    E' il controllo piu' importante di tutti: se il documento tornasse a
    prendere nome e IBAN dal template, chi usa l'app manderebbe fatture con
    l'IBAN di un altro e i clienti pagherebbero sul conto sbagliato. Un
    errore silenzioso, che si scopre solo quando i soldi non arrivano.
    """
    import tempfile
    from docx import Document
    from . import docgen, pdfgen
    from .db import DEFAULT_SETTINGS

    mio = dict(DEFAULT_SETTINGS,
               business_name='Anna Rossi Fitness', business_uid='CHE-111.222.333',
               business_addr1='Musterstrasse 1', business_addr2='Musterstadt, 8000',
               business_phone='+41 79 000 00 00', business_web='annarossi.ch',
               business_iban='CH5604835012345678009',
               terms='Payable within 10 days net to:')
    voci = [{'qty': 10, 'description': '10 Sessions Pack',
             'unit_cents': 12000, 'total_cents': 120000}]

    with tempfile.TemporaryDirectory() as tmp:
        dx = os.path.join(tmp, 'p.docx')
        pf = os.path.join(tmp, 'p.pdf')
        docgen.build_docx(dx, 7, '23-08-26', 'Mario Bianchi', ['Via Roma 1'], voci, 120000, mio)
        pdfgen.build_pdf(pf, 7, '23-08-26', 'Mario Bianchi', ['Via Roma 1'], voci, 120000, mio)
        d = Document(dx)
        testo = '\n'.join(p.text for p in d.paragraphs)
        for t in d.tables:
            for riga in t.rows:
                for c in riga.cells:
                    testo += '\n' + c.text
        for chiave, atteso in (('nome', 'Anna Rossi Fitness'), ('UID', 'CHE-111.222.333'),
                               ('indirizzo', 'Musterstrasse 1'), ('città', 'Musterstadt, 8000'),
                               ('telefono', '+41 79 000 00 00'), ('sito', 'annarossi.ch'),
                               ('IBAN', 'CH5604835012345678009'),
                               ('condizioni', 'Payable within 10 days net to:')):
            _check(r, 'Intestazione fattura', 'nel .docx c\'è il %s delle Impostazioni' % chiave,
                   atteso in testo, True)
        _check(r, 'Intestazione fattura', 'il ringraziamento nomina l\'attività',
               'Thanks for choosing Anna Rossi Fitness!' in testo, True)
        # se un segnaposto del template sopravvive vuol dire che quel dato non
        # e' stato scritto: meglio accorgersene qui che su una fattura spedita
        rimasti = [x for x in ('Nome attività', 'UID / IDI', 'Indirizzo, via e numero',
                               'CAP e città', 'IBAN: —', 'Nome cliente')
                   if x in testo]
        _check(r, 'Intestazione fattura', 'nel .docx non resta nessun segnaposto',
               rimasti, [])
        # Non basta guardare il testo: Word nasconde roba nei «controlli
        # contenuto» e nelle proprieta' del documento, e python-docx non la
        # mostra. Qui si apre il file come archivio e si guarda dentro tutto.
        import zipfile as _zip
        z = _zip.ZipFile(dx)
        dentro = '\n'.join(z.read(n).decode('utf-8', 'replace')
                           for n in z.namelist() if n.endswith('.xml'))
        # tutto quello che il template si porta dietro di chi l'ha disegnato:
        # non deve sopravvivere in nessuna fattura generata
        _check(r, 'Intestazione fattura', 'del template non resta niente nel file',
               [x for x in _proprieta_docx(_zip.ZipFile(docgen.TEMPLATE))
                if x and x in dentro], [])
        import re as _re
        # il ringraziamento era un campo agganciato alla proprieta' «Azienda»:
        # se resta agganciato, Word lo riempie da solo e il nome esce doppio
        corpo = z.read('word/document.xml').decode('utf-8', 'replace')
        _check(r, 'Intestazione fattura', 'niente campi agganciati alle proprietà',
               _re.findall(r'w:xpath="(.*?)"', corpo), [])
        _check(r, 'Intestazione fattura', 'il nome dell\'attività non esce doppio',
               corpo.count('Anna Rossi Fitness'), 3)
        core = z.read('docProps/core.xml').decode('utf-8', 'replace')
        app = z.read('docProps/app.xml').decode('utf-8', 'replace')
        for etichetta, trovato in (
                ('autore', _re.findall(r'<dc:creator>(.*?)</dc:creator>', core)),
                ('ultimo che ha scritto', _re.findall(r'<cp:lastModifiedBy>(.*?)</cp:lastModifiedBy>', core)),
                ('azienda', _re.findall(r'<Company>(.*?)</Company>', app))):
            _check(r, 'Intestazione fattura',
                   f'nelle proprietà del documento «{etichetta}» è la tua attività',
                   trovato, ['Anna Rossi Fitness'])
        _check(r, 'Intestazione fattura', 'gli importi restano giusti',
               verify.verify_generated(dx, pf, 120000, voci), [])
        # senza Impostazioni (chiamata vecchio stile) la fattura si fa lo stesso:
        # meglio un documento coi segnaposto che nessun documento
        dx2 = os.path.join(tmp, 'q.docx')
        docgen.build_docx(dx2, 8, '23-08-26', 'Mario Bianchi', ['Via Roma 1'], voci, 120000)
        _check(r, 'Intestazione fattura', 'senza Impostazioni non esplode',
               os.path.exists(dx2), True)


def _proprieta_docx(z):
    """Chi ha creato un .docx, chi l'ha modificato per ultimo, per quale azienda.

    Sono tre righe di XML che Word riempie da solo e che nessuno guarda mai:
    e' li' che il nome di chi ha disegnato il template resta appiccicato a ogni
    documento generato da quel template."""
    import re as _re
    fuori = []
    for parte, schema in (('docProps/core.xml', r'<dc:creator>(.*?)</dc:creator>'),
                          ('docProps/core.xml', r'<cp:lastModifiedBy>(.*?)</cp:lastModifiedBy>'),
                          ('docProps/app.xml', r'<Company>(.*?)</Company>')):
        try:
            testo = z.read(parte).decode('utf-8', 'replace')
        except KeyError:
            continue
        fuori += [v.strip() for v in _re.findall(schema, testo) if v.strip()]
    return fuori


def _test_marchio(r):
    """Il logo e il nome dell'attività vengono da chi usa l'app, non da chi
    l'ha scritta. Se il programma condiviso si portasse dietro il logo del
    primo proprietario, ogni utente manderebbe fatture col marchio di un altro."""
    import io as _io
    import tempfile
    import zipfile
    from PIL import Image
    from . import marchio, docgen
    from .db import DEFAULT_SETTINGS

    for nome, atteso in (('Studio Bianchi Fisioterapia', ('Studio Bianchi', 'Fisioterapia')),
                         ('Anna Rossi Personal Training', ('Anna Rossi', 'Personal Training')),
                         ('Centro Vitale', ('Centro Vitale', '')),
                         ('', ('La tua attività', ''))):
        _check(r, 'Marchio', 'il nome «%s» si spezza bene' % (nome or 'vuoto'),
               marchio.due_righe(nome), atteso)

    def _png(colore, misura=(120, 120)):
        buf = _io.BytesIO()
        Image.new('RGBA', misura, colore).save(buf, 'PNG')
        return buf.getvalue()

    vero = marchio.PERSONALE
    try:
        with tempfile.TemporaryDirectory() as tmp:
            marchio.PERSONALE = os.path.join(tmp, 'logo.png')
            _check(r, 'Marchio', 'senza logo caricato si usa il segnaposto',
                   marchio.percorso(), marchio.SEGNAPOSTO)
            _check(r, 'Marchio', 'il segnaposto esiste davvero',
                   os.path.exists(marchio.SEGNAPOSTO), True)
            # il segnaposto dice «caricalo dalle Impostazioni»: dentro l'app va
            # bene, su una fattura che parte al cliente sarebbe una figuraccia
            vuoto = Image.open(_io.BytesIO(marchio.adattato(60, 40)))
            _check(r, 'Marchio', 'senza logo la fattura lascia lo spazio vuoto',
                   vuoto.getbbox(), None)

            _check(r, 'Marchio', 'un file che non è un\'immagine viene rifiutato',
                   marchio.salva(b'questo non e\' un png') is not None, True)
            _check(r, 'Marchio', 'un caricamento vuoto viene rifiutato',
                   marchio.salva(b'') is not None, True)
            _check(r, 'Marchio', 'un\'immagine enorme viene rifiutata',
                   marchio.salva(b'\x89PNG' + b'x' * marchio.PESO_MAX) is not None, True)

            # un JPEG entra, ma quello che salviamo e' sempre un PNG: il resto
            # dell'app non deve sapere che formato aveva l'originale
            jpg = _io.BytesIO()
            Image.new('RGB', (900, 300), (200, 30, 30)).save(jpg, 'JPEG')
            _check(r, 'Marchio', 'un JPEG viene accettato', marchio.salva(jpg.getvalue()), None)
            _check(r, 'Marchio', 'adesso il logo è quello dell\'utente',
                   marchio.percorso(), marchio.PERSONALE)
            _check(r, 'Marchio', 'il logo salvato è un PNG',
                   Image.open(marchio.PERSONALE).format, 'PNG')
            _check(r, 'Marchio', 'un logo enorme viene rimpicciolito',
                   max(Image.open(marchio.PERSONALE).size) <= marchio.LATO_MAX, True)

            # dentro il Word lo spazio del logo ha una forma fissa: il logo ci
            # deve entrare con quella forma, ma senza essere stirato ne' rifatto
            marchio.salva(_png((0, 0, 255, 255), (300, 300)))
            fuori = Image.open(_io.BytesIO(marchio.adattato(200, 150)))
            _check(r, 'Marchio', 'il logo adattato prende la forma dello spazio',
                   round(fuori.width / fuori.height, 3), round(200 / 150, 3))
            _check(r, 'Marchio', 'un logo quadrato non viene schiacciato',
                   fuori.height, 300)
            # un logo gia' della forma giusta non va toccato per niente
            marchio.salva(_png((0, 0, 255, 255), (400, 300)))
            uguale = Image.open(_io.BytesIO(marchio.adattato(200, 150)))
            _check(r, 'Marchio', 'un logo già della forma giusta resta tale e quale',
                   uguale.size, (400, 300))

            marchio.salva(_png((255, 0, 0, 255)))
            with tempfile.TemporaryDirectory() as t2:
                dx = os.path.join(t2, 'p.docx')
                docgen.build_docx(dx, 9, '23-08-26', 'Mario Bianchi', [''],
                                  [{'qty': 1, 'description': 'x', 'unit_cents': 100,
                                    'total_cents': 100}], 100,
                                  dict(DEFAULT_SETTINGS, business_name='Anna Rossi Fitness'))
                dentro = Image.open(_io.BytesIO(
                    zipfile.ZipFile(dx).read('word/media/image1.png'))).convert('RGBA')
                rossi = [p for p in dentro.getdata() if p[3] > 0 and p[0] > 200 and p[1] < 60]
                _check(r, 'Marchio', 'nel .docx finisce il logo dell\'utente',
                       len(rossi) > 1000, True)

            marchio.rimuovi()
            _check(r, 'Marchio', 'tolto il logo si torna al segnaposto',
                   marchio.percorso(), marchio.SEGNAPOSTO)
            _check(r, 'Marchio', 'e la fattura torna a lasciare lo spazio vuoto',
                   Image.open(_io.BytesIO(marchio.adattato(60, 40))).getbbox(), None)
    finally:
        marchio.PERSONALE = vero

    # il template distribuito non deve contenere niente di nessuno: ne' il logo,
    # ne' il nome di chi l'ha disegnato nelle proprieta' del documento. Word ce
    # lo rimette ogni volta che qualcuno risalva il file.
    _check(r, 'Marchio', 'il template Word non porta il nome di chi l\'ha fatto',
           _proprieta_docx(zipfile.ZipFile(docgen.TEMPLATE)), [])
    interno = zipfile.ZipFile(docgen.TEMPLATE).read('word/media/image1.png')
    atteso = open(marchio.SEGNAPOSTO, 'rb').read()
    _check(r, 'Marchio', 'il template Word contiene solo il segnaposto',
           interno == atteso, True)


def _test_clienti_crediti(r):
    """Chi lavora a pacchetti si aggiunge dalla pagina Crediti, non nel codice.

    Qui si verifica che l'elenco scritto dall'utente comandi davvero tutto:
    quali nomi si riconoscono nel calendario, come si chiamano i pacchetti, a
    che prezzo una fattura viene riconosciuta, e la regola della coppia."""
    import tempfile
    from . import sessions as S

    for testo, atteso in ((' Giulia Ferrari ', 'giuliaferrari'), ('ANNA', 'anna'),
                          ('Jean-Luc', 'jeanluc'), ('', '')):
        _check(r, 'Clienti a crediti', f'«{testo}» diventa la parola «{atteso}»',
               S.chiave_da_nome(testo), atteso)
    for testo, atteso in (("1'800.00 CHF, 150.-", '180000,15000'), ('1800', '180000'),
                          ('2000; 2050', '200000,205000'), ('', ''), ('boh', '')):
        _check(r, 'Clienti a crediti', f'i prezzi «{testo}» diventano «{atteso}»',
               S.prezzi_da_testo(testo), atteso)

    prima = S._CONFIG
    try:
        S.configura([
            {'chiave': 'giulia', 'nome': 'Giulia', 'crediti': 10, 'prefisso': 'GIU',
             'prezzi': '150000', 'fattura_a': '', 'compagno': '', 'attivo': 1},
            {'chiave': 'marco', 'nome': 'Marco', 'crediti': 10, 'prefisso': '',
             'prezzi': '90000', 'fattura_a': 'Giulia', 'compagno': 'giulia', 'attivo': 1},
            {'chiave': 'luca', 'nome': 'Luca', 'crediti': 12, 'prefisso': 'LUC',
             'prezzi': '', 'fattura_a': '', 'compagno': '', 'attivo': 0},
        ])
        _check(r, 'Clienti a crediti', 'gli attivi sono due', sorted(S.clienti()), ['giulia', 'marco'])
        _check(r, 'Clienti a crediti', 'chi è archiviato sta a parte', S.ex_clienti(), {'luca': 'Luca'})
        _check(r, 'Clienti a crediti', 'il prezzo si legge come sulla fattura',
               S.prezzo_atteso('giulia'), "1'500.00 CHF")
        _check(r, 'Clienti a crediti', 'senza prezzo non si inventa niente',
               S.prezzo_atteso('luca'), None)

        _check(r, 'Clienti a crediti', 'nel calendario si riconosce chi è in elenco',
               S.classifica('Giulia')[0], 'giulia')
        _check(r, 'Clienti a crediti', 'si riconosce anche chi è archiviato',
               S.classifica('Luca')[0], 'luca')
        _check(r, 'Clienti a crediti', 'chi non è in elenco non conta',
               S.classifica('Federico')[0], None)
        _check(r, 'Clienti a crediti', 'il compleanno non è una sessione',
               S.classifica('Giulia birthday')[0], None)

        vuoto = {'pacchetti': [], 'esclusi': []}
        _check(r, 'Clienti a crediti', 'i pacchetti prendono la sigla scritta',
               S.prossimo_id_pacchetto(vuoto, 'giulia'), 'GIU-01')
        _check(r, 'Clienti a crediti', 'senza sigla se ne ricava una dal nome',
               S.prossimo_id_pacchetto(vuoto, 'marco'), 'MAR-01')

        # la regola della coppia: il supplemento vale solo se ci sono tutti e due
        _check(r, 'Clienti a crediti', 'in coppia il supplemento resta suo',
               S.attribuisci('marco', ['giulia', 'marco'])[0], 'marco')
        _check(r, 'Clienti a crediti', 'da solo consuma un credito del compagno',
               S.attribuisci('marco', ['marco'])[0], 'giulia')
        _check(r, 'Clienti a crediti', 'chi non è un supplemento non cambia mai',
               S.attribuisci('giulia', ['giulia'])[0], 'giulia')

        _check(r, 'Clienti a crediti', 'una fattura del prezzo giusto è un pacchetto',
               S.riconosci_pacchetto('Giulia Ferrari', 150000), 'giulia')
        _check(r, 'Clienti a crediti', 'un importo diverso non lo è',
               S.riconosci_pacchetto('Giulia Ferrari', 149900), None)
        # Marco e' fatturato a Giulia: senza «add-on» la fattura e' di Giulia
        _check(r, 'Clienti a crediti', 'un «add-on» va al supplemento',
               S.analizza_fattura('Giulia', 99900, ['Add-on 10 credits'])['chiave'], 'marco')
        _check(r, 'Clienti a crediti', 'un pacchetto normale va al cliente pieno',
               S.analizza_fattura('Giulia', 99900, ['10 Sessions Pack'])['chiave'], 'giulia')

        try:
            S.apri_pacchetto({'pacchetti': []}, 'federico', '2026-08-23')
            aperto = 'non ha protestato'
        except KeyError:
            aperto = 'protesta'
        _check(r, 'Clienti a crediti', 'non si apre un pacchetto a uno sconosciuto',
               aperto, 'protesta')

        # elenco vuoto: l'app deve reggere, non spegnersi
        S.configura([])
        _check(r, 'Clienti a crediti', 'senza nessun cliente la vista è vuota',
               S.vista_crediti({'pacchetti': []}), [])
        _check(r, 'Clienti a crediti', 'senza nessun cliente non si riconosce nulla',
               S.classifica('Giulia')[0], None)
    finally:
        S._CONFIG = prima

    # un'app appena installata non ha nessuno storico da ricopiare
    seed = S.SEED
    try:
        S.SEED = os.path.join(tempfile.gettempdir(), 'seed-che-non-esiste.json')
        with tempfile.TemporaryDirectory() as tmp:
            reg = S.carica(os.path.join(tmp, 'r.json'))
            _check(r, 'Clienti a crediti', 'senza storico il registro nasce vuoto',
                   (reg['pacchetti'], reg['esclusi']), ([], []))
    finally:
        S.SEED = seed


def _test_servizi(r):
    """I servizi proposti li scrive chi usa l'app, e il periodo degli
    abbonamenti lo sposta avanti l'app senza sapere come li hai chiamati."""
    from . import servizi as SR

    for testo, atteso in (
            ('Monthly abo: running coaching 13.07.26 – 12.08.26', 'Monthly abo: running coaching'),
            ('Abbonamento 01.01.2026 - 31.01.2026', 'Abbonamento'),
            ('10 Sessions Pack', '10 Sessions Pack')):
        _check(r, 'Servizi', f'«{testo}» senza il periodo', SR.senza_date(testo), atteso)

    for testo, atteso in (
            ('Monthly abo: running coaching 13.07.26 – 12.08.26',
             'Monthly abo: running coaching 13.08.26 – 12.09.26'),
            ('Abbonamento 01.01.2026 - 31.01.2026', 'Abbonamento 01.02.2026 - 28.02.2026'),
            ('Coaching 01/07/26 - 31/07/26', 'Coaching 01/08/26 - 31/08/26'),
            ('10 Sessions Pack', None),
            ('Coaching (August)', None),
            ('Coaching 32.13.26 - 40.99.26', None)):
        _check(r, 'Servizi', f'un mese avanti: «{testo}»', SR.avanza_periodo(testo), atteso)

    _check(r, 'Servizi', "l'anno a due cifre resta a due cifre",
           SR.avanza_periodo('X 01.12.26 - 31.12.26'), 'X 01.01.27 - 31.01.27')
    _check(r, 'Servizi', 'lo stesso servizio si riconosce nonostante il periodo',
           SR.stesso_servizio('Monthly abo', 'Monthly abo 01.07.26 - 31.07.26'), True)
    _check(r, 'Servizi', 'due servizi diversi non si confondono',
           SR.stesso_servizio('10 Sessions Pack', 'Monthly abo'), False)
    _check(r, 'Servizi', 'una descrizione vuota non somiglia a niente',
           SR.stesso_servizio('', 'Monthly abo'), False)

    # l'elenco: prima quello scritto a mano, altrimenti quello che si e' usato di piu'
    con = _db_finto()
    _check(r, 'Servizi', 'i servizi scritti a mano vincono',
           SR.elenco(con, {'servizi': 'Uno\nDue\n\n  Tre  '}), ['Uno', 'Due', 'Tre'])
    proposti = SR.elenco(con, {'servizi': ''})
    _check(r, 'Servizi', "senza elenco si propone quello che hai usato di piu'",
           proposti[0] if proposti else None, 'Monthly abo: running coaching')
    _check(r, 'Servizi', 'una descrizione usata una volta sola non si propone',
           'Una tantum' in proposti, False)
    con.close()


def _db_finto():
    """Un database in memoria con qualche fattura, per provare le proposte."""
    import sqlite3
    con = sqlite3.connect(':memory:')
    con.row_factory = sqlite3.Row
    con.executescript(
        'CREATE TABLE invoices(id INTEGER PRIMARY KEY, number INTEGER, deleted_at TEXT);'
        'CREATE TABLE items(id INTEGER PRIMARY KEY, invoice_id INTEGER, description TEXT);')
    righe = [('Monthly abo: running coaching 01.06.26 - 30.06.26',),
             ('Monthly abo: running coaching 01.07.26 - 31.07.26',),
             ('Monthly abo: running coaching 01.08.26 - 31.08.26',),
             ('10 Sessions Pack',), ('10 Sessions Pack',), ('Una tantum',)]
    for i, (d,) in enumerate(righe, 1):
        con.execute('INSERT INTO invoices(id, number, deleted_at) VALUES(?,?,NULL)', (i, i))
        con.execute('INSERT INTO items(invoice_id, description) VALUES(?,?)', (i, d))
    return con


def _test_lingua(r):
    """Le tre lingue.

    Il controllo che conta e' l'ultimo: ogni voce di menu deve esistere in
    tutte le lingue. Senza, il giorno che si aggiunge una pagina il menu esce
    metA' in italiano e meta' in tedesco, e nessuno se ne accorge finche' non
    lo vede un utente.
    """
    from . import lingua as L
    from . import menu as M

    _check(r, 'Lingua', 'le lingue sono tre', L.CODICI, ('it', 'en', 'de'))
    _check(r, 'Lingua', "un codice che non esiste torna all'italiano",
           L.normalizza('klingon'), 'it')
    _check(r, 'Lingua', 'senza lingua si resta in italiano',
           L.t('Fatture', None), 'Fatture')
    _check(r, 'Lingua', 'una frase mai tradotta resta in italiano invece di sparire',
           L.t('Frase che nessuno ha tradotto', 'de'), 'Frase che nessuno ha tradotto')
    _check(r, 'Lingua', 'tradurre in inglese funziona', L.t('Fatture', 'en'), 'Invoices')
    _check(r, 'Lingua', 'tradurre in tedesco funziona', L.t('Fatture', 'de'), 'Rechnungen')

    # nessuna voce di menu senza traduzione, in nessuna lingua
    da_tradurre = {etichetta for _t, voci in M.GRUPPI for _k, etichetta, _i, _a in voci}
    da_tradurre |= {titolo for titolo, _v in M.GRUPPI if titolo}
    da_tradurre.add(M.PRIMI_PASSI[1])
    for cod in ('en', 'de'):
        _check(r, 'Lingua', 'ogni voce del menu esiste in %s' % cod,
               sorted(f for f in da_tradurre if f not in L.TESTI[cod]), [])
    _check(r, 'Lingua', 'inglese e tedesco conoscono le stesse frasi',
           L.mancanti('en') + L.mancanti('de'), [])



def _db_fatture_finto(righe):
    """Un database in memoria con le sole colonne che il «Da fare» guarda."""
    import sqlite3
    con = sqlite3.connect(':memory:')
    con.row_factory = sqlite3.Row
    con.executescript(
        'CREATE TABLE invoices(id INTEGER PRIMARY KEY, number INTEGER, client_name TEXT,'
        ' date TEXT, total_cents INTEGER, status TEXT, paid_at TEXT, deleted_at TEXT,'
        ' year INTEGER, sent_at TEXT, source TEXT);'
        'CREATE TABLE settings(key TEXT PRIMARY KEY, value TEXT);')
    for i, r in enumerate(righe, 1):
        con.execute('INSERT INTO invoices(id, number, client_name, date, total_cents, status,'
                    ' paid_at, deleted_at, year, sent_at, source) VALUES(?,?,?,?,?,?,?,?,?,?,?)',
                    (i, i, r.get('cliente', 'X'), r['data'], r.get('cents', 10000),
                     r.get('stato', 'emessa'), r.get('paid_at'), r.get('deleted_at'),
                     int(r['data'][:4]), r.get('sent_at'), r.get('source', 'app')))
    return con


def _test_da_fare(r):
    """La lista «Da fare» della Dashboard.

    Prima la Dashboard mostrava nove riquadri dello stesso peso, e la cosa da
    fare andava cercata. Adesso e' una lista sola, e ogni voce e' un'azione. Il
    punto delicato e' che il numero deve corrispondere a quello che si vede
    cliccando: un contatore che dice 6 e apre una pagina con 3 righe fa perdere
    fiducia in tutti gli altri numeri della pagina.
    """
    import datetime
    from . import cruscotto as C

    oggi = datetime.date.today()
    ieri = (oggi - datetime.timedelta(days=1)).isoformat()
    vecchia = (oggi - datetime.timedelta(days=200)).isoformat()

    def voce(cose, chiave):
        return next((c for c in cose if c['chiave'] == chiave), None)

    # --- niente da fare: la lista e' vuota, e va bene cosi' ---
    con = _db_fatture_finto([{'data': ieri, 'stato': 'pagata', 'paid_at': ieri,
                              'sent_at': ieri}])
    vuoto = C.da_fare(con, {'banca_ultimo_estratto': oggi.isoformat()}, None)
    _check(r, 'Da fare', 'tutto a posto: nessuna voce', vuoto, [])

    # --- il numero e' quello delle fatture aperte, non quello della banca ---
    con = _db_fatture_finto([
        {'data': ieri, 'stato': 'emessa', 'cents': 100000, 'sent_at': ieri},
        {'data': ieri, 'stato': 'emessa', 'cents': 20000, 'sent_at': ieri},
        # pagata ma senza riscontro in banca: NON e' una cosa da fare
        {'data': ieri, 'stato': 'pagata', 'cents': 500000, 'sent_at': ieri},
    ])
    cose = C.da_fare(con, {'banca_ultimo_estratto': oggi.isoformat()}, None)
    inc = voce(cose, 'incassare')
    _check(r, 'Da fare', 'conta le fatture aperte, non quelle senza riscontro in banca',
           inc['quante'], 2)
    _check(r, 'Da fare', "l'importo e' la somma di quelle aperte", inc['importo'], 120000)
    _check(r, 'Da fare', 'il link porta esattamente a quelle contate',
           inc['link'], ('fatture', {'stato': 'emessa', 'anno': ''}))
    _check(r, 'Da fare', 'fatture recenti: in attesa, non in ritardo',
           inc['urgenza'], C.ATTESA)

    # --- una vecchia, con l'estratto aggiornato: quella e' in ritardo ---
    con = _db_fatture_finto([{'data': vecchia, 'stato': 'emessa', 'cents': 30000,
                              'sent_at': vecchia}])
    cose = C.da_fare(con, {'banca_ultimo_estratto': oggi.isoformat()}, None)
    _check(r, 'Da fare', 'ferma da mesi con estratto aggiornato: in ritardo',
           voce(cose, 'incassare')['urgenza'], C.RITARDO)

    # --- senza estratto non si accusa nessuno di ritardo ---
    con = _db_fatture_finto([{'data': vecchia, 'stato': 'emessa', 'cents': 30000,
                              'sent_at': vecchia}])
    cose = C.da_fare(con, {'banca_ultimo_estratto': ''}, None)
    _check(r, 'Da fare', 'senza estratto nessuna accusa di ritardo',
           voce(cose, 'incassare')['urgenza'], C.ATTESA)
    _check(r, 'Da fare', "senza estratto compare 'caricane uno'",
           voce(cose, 'estratto') is not None, True)

    # --- fatta con l'app e mai spedita ---
    con = _db_fatture_finto([{'data': oggi.isoformat(), 'stato': 'pagata',
                              'paid_at': oggi.isoformat(), 'source': 'app'}])
    cose = C.da_fare(con, {'banca_ultimo_estratto': oggi.isoformat()}, None)
    _check(r, 'Da fare', 'una fattura mai spedita si fa notare',
           voce(cose, 'spedire')['quante'], 1)
    # una storica importata non e' "da spedire": non e' mai partita dall'app
    con = _db_fatture_finto([{'data': oggi.isoformat(), 'stato': 'pagata',
                              'paid_at': oggi.isoformat(), 'source': 'import'}])
    cose = C.da_fare(con, {'banca_ultimo_estratto': oggi.isoformat()}, None)
    _check(r, 'Da fare', 'una fattura storica non conta come da spedire',
           voce(cose, 'spedire'), None)

    # --- l'estratto vecchio ---
    _check(r, 'Da fare', 'estratto di ieri: non si dice niente',
           C._estratto_vecchio(ieri), '')
    _check(r, 'Da fare', 'estratto di 200 giorni fa: si dice',
           '200 giorni fa' in C._estratto_vecchio(vecchia), True)
    _check(r, 'Da fare', 'una data storta non fa saltare la Dashboard',
           C._estratto_vecchio('non-una-data'), '')

    # --- i crediti non devono poter spegnere la Dashboard ---
    _check(r, 'Da fare', 'senza registro crediti, nessuna voce e nessun errore',
           C._crediti_finiti(None), [])
    _check(r, 'Da fare', 'un registro incomprensibile non fa saltare niente',
           C._crediti_finiti('non-un-registro'), [])



def _test_servizi_riconosciuti(r):
    """Quali servizi l'app riconosce nelle righe della fattura.

    Prima erano tre, scritti nel programma: quelli di chi l'app l'aveva
    scritta per se'. Chiunque altro si vedeva chiamare il proprio lavoro col
    nome del suo. Adesso le regole sono un'impostazione, e questi controlli
    servono a due cose: che chi le scrive ottenga quello che si aspetta, e che
    chi non le ha ancora scritte non si veda inventare niente.
    """
    from . import servizi as S
    from . import mailer, db

    mio = {'servizi_abbonamento': 'Running Coaching = running coaching\n'
                                  'Online Coaching = coaching online',
           'servizi_pacchetto': 'Personal Training = session, personal training'}

    # --- come si legge una riga ---
    _check(r, 'Servizi', 'la riga «Nome = parole» si legge tutta',
           S._regola('Fisioterapia = seduta, fisio'), ('Fisioterapia', ['seduta', 'fisio']))
    _check(r, 'Servizi', 'senza «=» il nome fa anche da parola',
           S._regola('Osteopatia'), ('Osteopatia', ['osteopatia']))
    _check(r, 'Servizi', 'una riga vuota non e\' una regola', S._regola('   '), None)
    _check(r, 'Servizi', 'gli spazi attorno alle parole non contano',
           S._regola('Massaggio =  a ,  b '), ('Massaggio', ['a', 'b']))

    # --- l'ordine: gli abbonamenti si provano per primi ---
    _check(r, 'Servizi', 'prima gli abbonamenti, poi i pacchetti',
           [n for n, _m, _p in S.regole(mio)],
           ['Running Coaching', 'Online Coaching', 'Personal Training'])
    _check(r, 'Servizi', 'ogni servizio porta con se\' il suo modello',
           [m for _n, m, _p in S.regole(mio)], ['coaching', 'coaching', 'pt'])

    # --- riconoscere ---
    _check(r, 'Servizi', 'riconosce l\'abbonamento e sa che e\' un abbonamento',
           S.riconosci('Monthly abo: running coaching (August)', mio),
           ('Running Coaching', 'coaching'))
    _check(r, 'Servizi', 'riconosce il pacchetto',
           S.riconosci('10 Sessions Pack', mio), ('Personal Training', 'pt'))
    _check(r, 'Servizi', 'le maiuscole non contano',
           S.riconosci('COACHING ONLINE - agosto', mio), ('Online Coaching', 'coaching'))
    _check(r, 'Servizi', 'quello che non e\' scritto non viene riconosciuto',
           S.riconosci('Consulenza nutrizionale', mio), (None, None))
    _check(r, 'Servizi', 'senza regole scritte non si indovina niente',
           S.riconosci('Monthly abo: running coaching', {}), (None, None))
    _check(r, 'Servizi', 'una riga vuota non riconosce niente',
           S.riconosci('', mio), (None, None))

    # --- un altro mestiere, con le sue parole ---
    fisio = {'servizi_abbonamento': 'Riabilitazione = riabilitazione',
             'servizi_pacchetto': 'Fisioterapia = seduta, sedute, fisioterapia'}
    _check(r, 'Servizi', 'un fisioterapista riconosce le proprie righe',
           S.riconosci('Pacchetto 10 sedute di fisioterapia', fisio),
           ('Fisioterapia', 'pt'))
    _check(r, 'Servizi', 'e le sue righe non diventano quelle di un altro',
           S.riconosci('Pacchetto 10 sedute di fisioterapia', mio), (None, None))

    # --- e l'email che ne esce ---
    inv = _Finta(number=99, total_cents=110000, pdf_path='', source_file='',
                 client_name='Sofia Ferrari')
    cli = _Finta(name='Sofia Ferrari', email='s@esempio.ch', abbonamento=0, tono='informale')
    imp = dict(db.DEFAULT_SETTINGS, **fisio)
    corpo = mailer.componi(inv, cli, imp, ['Pacchetto 10 sedute di fisioterapia'])['body']
    _check(r, 'Servizi', 'il fisioterapista fattura a nome suo, non di un altro',
           'Please find attached your invoice for Fisioterapia.' in corpo, True)
    vuote = mailer.componi(inv, cli, dict(db.DEFAULT_SETTINGS),
                           ['Pacchetto 10 sedute di fisioterapia'])['body']
    _check(r, 'Servizi', 'senza regole l\'email non nomina nessun servizio',
           'Please find attached your invoice.' in vuote, True)



def _test_primi_passi(r):
    """Un'app appena installata deve dire da dove si comincia, e smettere di
    dirlo appena non serve piu'."""
    import sqlite3
    import tempfile
    from . import benvenuto as B, marchio
    from .db import DEFAULT_SETTINGS

    # il passo del logo guarda il file vero: qui si guarda altrove, altrimenti
    # il controllo dipende da chi lo sta eseguendo
    logo_vero = marchio.PERSONALE
    marchio.PERSONALE = os.path.join(tempfile.gettempdir(), 'logo-che-non-esiste.png')
    con = sqlite3.connect(':memory:')
    con.row_factory = sqlite3.Row
    con.executescript(
        'CREATE TABLE clients(id INTEGER PRIMARY KEY, archived INTEGER DEFAULT 0);'
        'CREATE TABLE invoices(id INTEGER PRIMARY KEY, deleted_at TEXT);'
        'CREATE TABLE crediti_clienti(chiave TEXT PRIMARY KEY);')

    vuoto = B.passi(con, dict(DEFAULT_SETTINGS))
    _check(r, 'Primi passi', 'appena installata non è fatto niente',
           B.avanzamento(vuoto)[0], 0)
    _check(r, 'Primi passi', 'appena installata manca l\'essenziale',
           B.manca_l_essenziale(vuoto), True)
    _check(r, 'Primi passi', 'ogni passo sa dove mandarti',
           [p['chiave'] for p in vuoto if not p['dove']], [])

    # i dati dell'attivita' senza IBAN non bastano: la fattura uscirebbe senza
    # il conto su cui incassare
    mezzo = dict(DEFAULT_SETTINGS, business_name='Studio Bianchi',
                 business_addr1='Via Roma 1', business_addr2='6900 Lugano')
    _check(r, 'Primi passi', "senza IBAN manca ancora l'essenziale",
           B.manca_l_essenziale(B.passi(con, mezzo)), True)

    pieno = dict(mezzo, business_iban='CH5604835012345678009')
    passi = B.passi(con, pieno)
    _check(r, 'Primi passi', "con nome, indirizzo e IBAN l'essenziale c'è",
           B.manca_l_essenziale(passi), False)
    _check(r, 'Primi passi', 'ma resta ancora qualcosa da fare',
           len(B.da_fare(passi)) > 0, True)

    # le cose non essenziali si spuntano da sole quando succedono
    con.execute('INSERT INTO clients(id, archived) VALUES(1, 0)')
    con.execute('INSERT INTO invoices(id, deleted_at) VALUES(1, NULL)')
    con.execute("INSERT INTO crediti_clienti(chiave) VALUES('anna')")
    dopo = {p['chiave']: p['fatto'] for p in B.passi(con, pieno)}
    for chiave in ('clienti', 'fattura', 'crediti'):
        _check(r, 'Primi passi', f'«{chiave}» si spunta da solo', dopo[chiave], True)
    # un cliente archiviato non conta come cliente
    con.execute('UPDATE clients SET archived = 1')
    _check(r, 'Primi passi', 'un cliente archiviato non conta',
           {p['chiave']: p['fatto'] for p in B.passi(con, pieno)}['clienti'], False)
    marchio.PERSONALE = logo_vero
    con.close()


def _sorgenti(cartella, estensione):
    """Il testo di tutti i file di un tipo, con il nome davanti."""
    base = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), cartella)
    fuori = {}
    for nome in sorted(os.listdir(base)):
        if nome.endswith(estensione):
            with io.open(os.path.join(base, nome), encoding='utf-8') as f:
                fuori[nome] = f.read()
    return fuori


def _test_icone(r):
    """I disegni al posto delle emoji.

    Il rischio di un'icona sbagliata non e' che si veda male: e' che non si
    veda affatto, perche' icona() con un nome che non esiste non disegna
    niente pur di non far saltare la pagina. Qui si controlla che tutti i nomi
    usati nelle pagine esistano davvero.
    """
    import xml.etree.ElementTree as ET
    from . import icone as I

    modelli = _sorgenti('templates', '.html')
    tutto = '\n'.join(modelli.values())

    # ogni disegno e' un pezzo di XML valido, dentro la stessa cornice
    rotti, fuori_cornice = [], []
    for nome in I.nomi():
        try:
            ET.fromstring('<svg>%s</svg>' % I.DISEGNI[nome])
        except Exception:
            rotti.append(nome)
        marcatura = str(I.icona(nome))
        if 'viewBox="0 0 24 24"' not in marcatura or 'stroke="currentColor"' not in marcatura:
            fuori_cornice.append(nome)
    _check(r, 'Icone', 'tutti i disegni sono XML valido', rotti, [])
    _check(r, 'Icone', 'tutti stanno nello stesso quadrato e prendono il colore del testo',
           fuori_cornice, [])

    # ogni nome usato nelle pagine esiste
    usati = set(re.findall(r"icona\(\s*'([a-z_]+)'", tutto))
    usati |= {v[2] for v in menu.voci()}
    usati |= set(re.findall(r"'(?:fattura|email|errore|sessione|crediti)':\s*'([a-z_]+)'", tutto))
    _check(r, 'Icone', 'nessuna pagina chiede un disegno che non esiste',
           sorted(usati - set(I.nomi())), [])
    _check(r, 'Icone', 'ci sono disegni da usare', len(usati) > 15, True)

    # niente disegni tenuti da parte «per dopo»
    citati = {n for n in I.nomi() if ("'%s'" % n) in tutto or n in str(menu.GRUPPI)}
    _check(r, 'Icone', 'nessun disegno rimasto inutilizzato',
           sorted(set(I.nomi()) - citati), [])

    # un nome sbagliato non fa saltare la pagina
    _check(r, 'Icone', 'un nome inventato non rompe la pagina', str(I.icona('nonesiste')), '')

    # i pallini di stato
    _check(r, 'Icone', 'il pallino verde ha la sua classe',
           'class="pallino verde"' in str(I.pallino('verde')), True)
    _check(r, 'Icone', 'il pallino spiega cosa vuol dire',
           'title="ha pagato"' in str(I.pallino('verde', 'ha pagato')), True)
    _check(r, 'Icone', 'un colore inventato ripiega sul pallino vuoto',
           'pallino vuoto' in str(I.pallino('fucsia')), True)
    _check(r, 'Icone', 'il titolo del pallino non puo\' iniettare marcatura',
           '<b>' in str(I.pallino('verde', '<b>x</b>')), False)

    # nelle pagine non devono restare emoji colorate
    faccine = re.compile('[\U0001F300-\U0001FAFF\U0001F004-\U0001F0CF]')
    # la pagina dell'errore tiene la sua faccina: li' non e' un'icona, e' il
    # tono con cui l'app si scusa, ed e' l'unico punto dove serve una faccia
    with_emoji = sorted(n for n, t in modelli.items()
                        if n != 'errore.html' and faccine.search(t))
    _check(r, 'Icone', 'nessuna emoji colorata rimasta nelle pagine', with_emoji, [])
    _check(r, 'Icone', "la faccina resta solo nella pagina dell'errore",
           bool(faccine.search(modelli['errore.html'])), True)
    js = _sorgenti('static', '.js')
    _check(r, 'Icone', 'nessuna emoji colorata rimasta nel codice delle pagine',
           sorted(n for n, t in js.items() if faccine.search(t)), [])


def _test_menu(r):
    """Il menu di sinistra: gruppi, ordine e voce accesa."""
    from . import menu as M

    endpoint = [v[0] for v in M.voci()]
    _check(r, 'Menu', 'nessuna voce ripetuta', len(endpoint), len(set(endpoint)))

    # ogni voce punta a una pagina che esiste davvero
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    with io.open(os.path.join(base, 'app.py'), encoding='utf-8') as f:
        programma = f.read()
    pagine = set(re.findall(r'^def ([a-z_0-9]+)\(', programma, re.M))
    _check(r, 'Menu', 'ogni voce porta a una pagina che esiste',
           sorted(set(endpoint) - pagine), [])

    # una pagina non puo' accendere due voci insieme
    acceso = {}
    doppie = []
    for e, _etichetta, _disegno, attivo in M.voci():
        for pagina in attivo:
            if pagina in acceso:
                doppie.append(pagina)
            acceso[pagina] = e
    _check(r, 'Menu', 'nessuna pagina accende due voci insieme', sorted(doppie), [])
    _check(r, 'Menu', 'ogni voce accende almeno se stessa',
           [e for e, _t, _d, a in M.voci() if e not in a], [])

    # le pagine di dettaglio devono accendere la voce del loro elenco:
    # senza, aprendo una fattura il menu si spegne e non si capisce dove si e'
    _check(r, 'Menu', 'la scheda di una fattura accende «Fatture»',
           acceso.get('fattura'), 'fatture')
    _check(r, 'Menu', "l'email di una fattura accende «Fatture»",
           acceso.get('fattura_email'), 'fatture')
    _check(r, 'Menu', 'una mail letta accende «Email inviate»',
           acceso.get('email_letta'), 'email_inviate')
    _check(r, 'Menu', 'un pacchetto accende «Crediti»',
           acceso.get('crediti_pacchetto'), 'crediti')
    _check(r, 'Menu', 'chi lavora a crediti accende «Crediti»',
           acceso.get('crediti_clienti'), 'crediti')

    _check(r, 'Menu', 'le voci stanno in gruppi con un titolo',
           [t for t, _ in M.GRUPPI if t], ['Fatturare', 'Chi segui', 'Incassi e fisco', "L'app"])
    _check(r, 'Menu', 'la Dashboard sta in cima, fuori dai gruppi',
           M.GRUPPI[0][0] is None and M.GRUPPI[0][1][0][0] == 'dashboard', True)
    _check(r, 'Menu', 'i primi passi stanno fuori dai gruppi fissi',
           M.PRIMI_PASSI[0], 'benvenuto')


def _test_finestra_stretta(r):
    """Cosa deve reggere quando la finestra e' stretta.

    L'app non si apre dal telefono, ma la finestra si tiene volentieri a meta'
    schermo. Il guaio non e' che diventi brutta: e' che la pagina scorra di
    lato, perche' allora il menu scompare e per leggere l'ultima colonna devi
    trascinare tutto. Qui si controllano le tre cose da cui dipende.
    """
    base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    with io.open(os.path.join(base, 'static', 'style.css'), encoding='utf-8') as f:
        stile = f.read()

    def regola(selettore):
        m = re.search(re.escape(selettore) + r'\s*\{(.*?)\}', stile, re.S)
        return m.group(1) if m else ''

    # 1. senza min-width:0 un elemento flessibile non si stringe sotto la
    #    larghezza del suo contenuto, e una tabella larga allarga la pagina
    _check(r, 'Finestra stretta', 'la parte centrale puo\' stringersi',
           'min-width: 0' in regola('main'), True)

    # 2. il contenuto troppo largo deve scorrere dentro il suo riquadro
    _check(r, 'Finestra stretta', 'i riquadri fanno scorrere dentro di se\'',
           'overflow-x: auto' in regola('.panel'), True)

    # 3. ...il che serve a qualcosa solo se le tabelle stanno nei riquadri
    modelli = _sorgenti('templates', '.html')
    sfuse = {n: _tabelle_fuori_dai_riquadri(t) for n, t in modelli.items()}
    _check(r, 'Finestra stretta', 'ogni tabella sta dentro un riquadro',
           sorted(n for n, q in sfuse.items() if q), [])
    _check(r, 'Finestra stretta', 'di tabelle ce ne sono, non e\' un controllo a vuoto',
           sum(t.count('<table') for t in modelli.values()) > 20, True)

    # le soglie: dalla piu' larga alla piu' stretta, senza buchi
    soglie = [int(x) for x in re.findall(r'@media \(max-width: (\d+)px\)', stile)]
    _check(r, 'Finestra stretta', 'le soglie sono in ordine dalla piu\' larga',
           soglie, sorted(soglie, reverse=True))
    _check(r, 'Finestra stretta', 'il menu diventa una colonnina di icone',
           1040 in soglie, True)


def _tabelle_fuori_dai_riquadri(testo):
    """Quante tabelle non hanno un «panel» sopra di loro."""
    senza_jinja = re.sub(r'\{[%{].*?[%}]\}', '', testo, flags=re.S)
    pila, fuori = [], 0
    for m in re.finditer(r'<(/?)(div|table)\b([^>]*)>', senza_jinja):
        chiusura, tag, attr = m.groups()
        if tag == 'table':
            if not chiusura and not any('panel' in c for c in pila):
                fuori += 1
        elif chiusura:
            if pila:
                pila.pop()
        elif not attr.rstrip().endswith('/'):
            cls = re.search(r'class="([^"]*)"', attr)
            pila.append(cls.group(1) if cls else '')
    return fuori


def _test_calendario(r):
    """L'iCal descrive le serie in modo compatto: espanderle bene e' delicato."""
    import datetime
    from . import calendario as C

    # il nome del calendario lo dice il calendario: cosi' nel programma non
    # resta scritto come si chiama quello di nessuno
    _check(r, 'Calendario', 'il calendario dice come si chiama', C.nome(CAL_PROVA),
           'Allenamenti')
    _check(r, 'Calendario', 'un calendario senza nome non ne inventa uno',
           C.nome('BEGIN:VCALENDAR\nEND:VCALENDAR'), '')
    _check(r, 'Calendario', 'il nome si legge anche se la riga e\' spezzata in due',
           C.nome('BEGIN:VCALENDAR\nX-WR-CALNAME:Alle\n namenti PT\nEND:VCALENDAR'),
           'Allenamenti PT')

    ev = C.leggi(CAL_PROVA, datetime.date(2026, 8, 20), datetime.date(2026, 9, 10),
                 e_testo=True)
    date = lambda t: sorted(e['data'] for e in ev if e['titolo'].startswith(t))

    _check(r, 'Calendario', 'i martedì di Marco si espandono',
           date('Marco'), ['2026-08-25', '2026-09-08'])
    _check(r, 'Calendario', 'la data in EXDATE (01.09) non compare',
           '2026-09-01' in date('Marco'), False)
    _check(r, 'Calendario', 'la sessione disdetta resta ma marcata cancellata',
           [e['stato_google'] for e in ev if e['data'] == '2026-08-27'], ['cancelled'])
    _check(r, 'Calendario', 'la virgola scritta come \\, torna normale',
           any(', spostata' in e['titolo'] for e in ev), True)
    _check(r, 'Calendario', "l'identificativo è UID + data, non solo UID",
           [e['id'] for e in ev if e['data'] == '2026-08-25'], ['tizio@g::2026-08-25'])
    _check(r, 'Calendario', 'rileggere lo stesso calendario dà gli stessi id '
           '(niente doppi conteggi)',
           [e['id'] for e in C.leggi(CAL_PROVA, datetime.date(2026, 8, 20),
                                     datetime.date(2026, 9, 10), e_testo=True)],
           [e['id'] for e in ev])
    _check(r, 'Calendario', 'fuori finestra non si legge niente',
           C.leggi(CAL_PROVA, datetime.date(2026, 1, 1), datetime.date(2026, 1, 31),
                   e_testo=True), [])
    _check(r, 'Calendario', 'un file illeggibile non fa esplodere niente',
           C.leggi('roba a caso', datetime.date(2026, 8, 20),
                   datetime.date(2026, 8, 31), e_testo=True), [])

    # --- l'ora, che serve all'Agenda ---
    _check(r, 'Calendario', "l'ora di inizio si legge dalla serie",
           [e['ora'] for e in ev if e['data'] == '2026-08-25'], ['07:30'])
    _check(r, 'Calendario', "ogni ripetizione eredita l'ora della serie",
           sorted({e['ora'] for e in ev if e['titolo'].startswith('Marco')}), ['07:30'])
    _check(r, 'Calendario', 'ora scritta in UTC riportata a quella svizzera',
           C._ora('20260825T053000Z'), '07:30')
    _check(r, 'Calendario', 'ora con fuso esplicito presa così com\'è',
           C._ora('20260825T073000'), '07:30')
    _check(r, 'Calendario', 'evento di sola giornata: nessuna ora inventata',
           C._ora('20260825'), None)

    _test_agenda(r)


def _test_agenda(r):
    """L'agenda mette insieme due fonti: registro (quali) e calendario (a che ora)."""
    from . import agenda as A

    reg = {'pacchetti': [
        {'id': 'X-01', 'cliente': 'Tizia', 'crediti': 10, 'fattura_numero': 7,
         'sessioni': [
             {'n': 1, 'data': '2026-08-18', 'titolo': 'Tizia', 'cancellata': False},
             {'n': 2, 'data': '2026-08-20', 'titolo': 'Tizia', 'cancellata': True},
             {'n': 3, 'data': '2026-08-20', 'titolo': 'Tizia', 'cancellata': False,
              'ora': '18:00'},
         ]},
        {'id': 'Y-01', 'cliente': 'Caio', 'crediti': 12, 'sessioni': [
            {'n': 1, 'data': '2025-03-01', 'titolo': 'Caio', 'cancellata': False}]},
    ]}
    orari = {A._chiave('2026-08-18', 'Tizia'): '07:30'}
    righe = A.elenco(reg, orari)

    _check(r, 'Agenda', 'ci sono tutte le sessioni del registro', len(righe), 4)
    _check(r, 'Agenda', 'la più recente sta in cima', righe[0]['data'], '2026-08-20')
    _check(r, 'Agenda', "l'ora arriva dall'indice del calendario",
           [x['ora'] for x in righe if x['data'] == '2026-08-18'], ['07:30'])
    _check(r, 'Agenda', "l'ora scritta nella sessione batte l'indice",
           righe[0]['ora'], '18:00')
    _check(r, 'Agenda', 'senza ora non se ne inventa una',
           [x['ora'] for x in righe if x['data'] == '2025-03-01'], [None])
    _check(r, 'Agenda', 'la sessione annullata resta (ha consumato il credito)',
           sum(1 for x in righe if x['cancellata']), 1)
    _check(r, 'Agenda', 'il numero di fattura del pacchetto arriva in riga',
           righe[0]['fattura'], 7)
    _check(r, 'Agenda', 'filtro per cliente',
           [x['cliente'] for x in A.elenco(reg, orari, cliente='caio')], ['Caio'])
    _check(r, 'Agenda', 'filtro per anno',
           len(A.elenco(reg, orari, anno='2026')), 3)
    _check(r, 'Agenda', 'gli anni disponibili sono quelli veri',
           A.anni(reg), ['2026', '2025'])
    _check(r, 'Agenda', 'il riepilogo conta le sessioni con orario',
           A.riepilogo(righe)['con_ora'], 2)

    _test_registro_email(r)


def _test_registro_email(r):
    """Il registro delle email si prova su un database usa e getta."""
    import sqlite3
    from .db import SCHEMA

    con = sqlite3.connect(':memory:')
    con.row_factory = sqlite3.Row
    con.executescript(SCHEMA)
    ins = ('INSERT INTO email_log(sent_at, destinatario, fatture, prova, esito, motivo, '
           'ccn, corpo) VALUES(?,?,?,?,?,?,?,?)')
    con.execute(ins, ('2026-08-21T10:00:00', 'a@b.ch', '84', 0, 'ok', '', 'io@b.ch',
                      'Dear X,\n\nriga due\n'))
    con.execute(ins, ('2026-08-21T11:00:00', 'io@b.ch', '84', 1, 'ok', '', '', 'prova'))
    con.execute(ins, ('2026-08-21T12:00:00', 'a@b.ch', '85', 0, 'errore', 'timeout', '', ''))

    conta = lambda q: con.execute('SELECT COUNT(*) c FROM email_log' + q).fetchone()['c']
    _check(r, 'Email inviate', 'la tabella esiste e accetta le righe', conta(''), 3)
    _check(r, 'Email inviate', 'le prove si separano da quelle ai clienti',
           conta(' WHERE prova=0'), 2)
    _check(r, 'Email inviate', 'i tentativi falliti restano scritti',
           conta(" WHERE esito='errore'"), 1)
    _check(r, 'Email inviate', 'la più recente viene per prima',
           con.execute('SELECT sent_at FROM email_log ORDER BY sent_at DESC, id DESC '
                       'LIMIT 1').fetchone()['sent_at'], '2026-08-21T12:00:00')
    prima = con.execute('SELECT * FROM email_log ORDER BY id LIMIT 1').fetchone()
    _check(r, 'Email inviate', 'il testo si rilegge tale e quale, a capo compresi',
           prima['corpo'], 'Dear X,\n\nriga due\n')
    _check(r, 'Email inviate', 'la copia nascosta resta scritta', prima['ccn'], 'io@b.ch')
    con.close()

    # le intestazioni salvate sono quelle vere, e la prova non ha la copia nascosta
    import app as APP
    from .db import DEFAULT_SETTINGS as S
    msg = {'body': 'ciao'}
    i = APP._intestazioni(msg, dict(S, email_copia_a_me='1'), None)
    _check(r, 'Email inviate', "l'invio al cliente registra la copia nascosta",
           i['ccn'], S['smtp_user'])
    i2 = APP._intestazioni(msg, dict(S, email_copia_a_me='1'), 'io@esempio.ch')
    _check(r, 'Email inviate', 'la prova a te stesso non ha copia nascosta da registrare',
           i2['ccn'], '')
    i3 = APP._intestazioni(msg, dict(S, email_copia_a_me='0'), None)
    _check(r, 'Email inviate', 'copia nascosta spenta: non si scrive niente', i3['ccn'], '')

    _test_cruscotto(r)


def _test_cruscotto(r):
    """I tre riquadri della Dashboard: contano e ordinano, non inventano."""
    import sqlite3
    import datetime
    from . import cruscotto as C
    from . import db as _db
    from .db import SCHEMA, DEFAULT_SETTINGS as S

    con = sqlite3.connect(':memory:')
    con.row_factory = sqlite3.Row
    con.executescript(SCHEMA)
    _db._migrate(con)          # le colonne aggiunte dopo (sent_at, email...) servono qui
    ins = ('INSERT INTO invoices(number, client_name, date, year, total_cents, status, '
           'source, created_at, sent_at) VALUES(?,?,?,?,?,?,?,?,?)')
    con.execute(ins, (1, 'A', '2026-01-10', 2026, 100, 'pagata', 'app',
                      '2026-01-10T09:00:00', '2026-01-10T10:00:00'))
    con.execute(ins, (2, 'B', '2026-02-10', 2026, 100, 'emessa', 'app',
                      '2026-02-10T09:00:00', None))
    con.execute(ins, (3, 'C', '2026-03-10', 2026, 100, 'emessa', 'storico',
                      '2026-03-10T09:00:00', None))
    con.execute(ins, (4, 'D', '2025-03-10', 2025, 100, 'pagata', 'app',
                      '2025-03-10T09:00:00', None))
    con.commit()

    st = C.stato_fatture(con, 2026)
    _check(r, 'Cruscotto', "l'anno chiesto è l'unico contato", st['totali'], 3)
    _check(r, 'Cruscotto', 'pagate e da incassare fanno il totale',
           st['pagate'] + st['da_incassare'], st['totali'])
    _check(r, 'Cruscotto', 'spedite: solo quelle con la data di invio', st['inviate'], 1)
    _check(r, 'Cruscotto', 'da mandare: le storiche non si contano (non si spediscono)',
           st['da_mandare'], 1)

    # salute: senza cartella di destinazione il backup è un problema, non un dettaglio
    sal = C.salute(con, dict(S, smtp_pass='', calendario_ics=''),
                   '/questa/cartella/non/esiste')
    voce = next(v for v in sal['voci'] if v['nome'] == 'Copia fuori dal Mac')
    _check(r, 'Cruscotto', 'backup mancante: rosso, non verde', voce['stato'], C.ROSSO)
    _check(r, 'Cruscotto', 'basta una voce rossa perché il riquadro lo dica',
           sal['stato'], C.ROSSO)

    reg = {'pacchetti': [{'id': 'X-01', 'cliente': 'Tizia', 'crediti': 10, 'rimasti': 0,
                          'fine': '2026-02-20', 'sessioni': [
                              {'n': 1, 'data': '2026-02-20', 'titolo': 'Tizia',
                               'cancellata': False, 'ora': '18:30'}]}]}
    voci = C.attivita(con, reg, quante=10)
    _check(r, 'Cruscotto', 'le fonti diverse finiscono in ordine di tempo',
           [v['quando'][:10] for v in voci], sorted([v['quando'][:10] for v in voci],
                                                    reverse=True))
    sess_ = next(v for v in voci if v['tipo'] == 'sessione')
    _check(r, 'Cruscotto', "l'ora vera della sessione arriva nel diario",
           sess_['quando'][11:16], '18:30')
    _check(r, 'Cruscotto', "il pacchetto finito compare senza un'ora inventata",
           next(v['ora_nota'] for v in voci if v['tipo'] == 'crediti'), False)
    _check(r, 'Cruscotto', 'un tempo di un\'ora si dice al singolare',
           C._eta((datetime.datetime.now() - datetime.timedelta(hours=1)).isoformat())[1],
           "un'ora fa")
    con.close()

    _test_incassi(r)
    _test_banca(r)


CSV_PROVA = """Estratto conto;;;
Conto;CH93 0076 2011 6238 5295 7;;
;;;
Data contabile;Data valuta;Descrizione;Accredito;Addebito
02.08.2026;02.08.2026;E-Banking Auftrag Sofia Ferrari;110.00;
05.08.2026;05.08.2026;LSV Krankenkasse;;432.10
14.08.2026;14.08.2026;Zahlung Mueller Petra;2'000.00;
20.08.2026;20.08.2026;Gutschrift unbekannt;110.00;
"""

CAMT_PROVA = """<?xml version="1.0" encoding="UTF-8"?>
<Document xmlns="urn:iso:std:iso:20022:tech:xsd:camt.054.001.04"><BkToCstmrDbtCdtNtfctn><Ntfctn>
 <Ntry><Amt Ccy="CHF">1800.00</Amt><CdtDbtInd>CRDT</CdtDbtInd>
  <BookgDt><Dt>2026-06-05</Dt></BookgDt>
  <NtryDtls><TxDtls><RltdPties><Dbtr><Nm>Bruno Keller</Nm></Dbtr></RltdPties></TxDtls></NtryDtls>
 </Ntry>
 <Ntry><Amt Ccy="CHF">55.00</Amt><CdtDbtInd>DBIT</CdtDbtInd>
  <BookgDt><Dt>2026-06-06</Dt></BookgDt></Ntry>
</Ntfctn></BkToCstmrDbtCdtNtfctn></Document>
"""


def _test_incassi(r):
    """La differenza fra «non pagata» e «non ancora verificabile»."""
    import sqlite3
    from . import db as _db
    from . import cruscotto as C
    from .db import SCHEMA

    con = sqlite3.connect(':memory:')
    con.row_factory = sqlite3.Row
    con.executescript(SCHEMA)
    _db._migrate(con)
    ins = ('INSERT INTO invoices(number, client_name, date, year, total_cents, status, '
           'source) VALUES(?,?,?,?,?,?,"app")')
    con.execute(ins, (10, 'Tizia', '2026-01-10', 2026, 11000, 'emessa'))   # vecchia, scoperta
    con.execute(ins, (11, 'Tizia', '2026-07-25', 2026, 11000, 'emessa'))   # troppo recente
    con.execute(ins, (12, 'Tizia', '2026-08-20', 2026, 11000, 'emessa'))   # dopo l'estratto
    con.commit()

    m = C.incassi_mancanti(con, '2026-07-31')
    _check(r, 'Incassi', 'una fattura vecchia e scoperta è un ritardo vero',
           [x['number'] for x in m['in_ritardo']], [10])
    _check(r, 'Incassi', 'una fattura recente non è in ritardo, è solo da verificare',
           sorted(x['number'] for x in m['da_verificare']), [11, 12])
    _check(r, 'Incassi', 'il confine sono 45 giorni prima dell\'ultimo estratto',
           m['limite'], '2026-06-16')

    con.execute('UPDATE invoices SET paid_at="2026-02-01" WHERE number=10')
    con.commit()
    _check(r, 'Incassi', 'incassata: sparisce dai ritardi',
           C.incassi_mancanti(con, '2026-07-31')['in_ritardo'], [])

    senza = C.incassi_mancanti(con, '')
    _check(r, 'Incassi', 'senza estratti non si accusa nessuno di ritardo',
           senza['in_ritardo'], [])
    con.close()


def _test_intestatario(r):
    """Chi fa le sedute e chi riceve la fattura possono essere due persone diverse."""
    import sqlite3
    from . import db as _db
    from .db import SCHEMA

    con = sqlite3.connect(':memory:')
    con.row_factory = sqlite3.Row
    con.executescript(SCHEMA)
    _db._migrate(con)
    con.execute("INSERT INTO clients(key, name, address1, address2, file_label, intestatario) "
                "VALUES('anna','Anna','Via delle Prove 1','8000 Cittanova','Anna','Luca Conti')")
    con.execute("INSERT INTO clients(key, name, address1, address2, file_label) "
                "VALUES('sara','Sara Bernasconi','Via delle Prove 2','8000 Cittanova','Sara B')")
    con.commit()

    def intestata_a(chiave):
        c = con.execute('SELECT * FROM clients WHERE key=?', (chiave,)).fetchone()
        return (c['intestatario'] or '').strip() or c['name']

    def nome_file(chiave):
        c = con.execute('SELECT * FROM clients WHERE key=?', (chiave,)).fetchone()
        return (c['file_label'] or c['name']).strip()

    _check(r, 'Intestatario', 'con intestatario: la fattura va a lui',
           intestata_a('anna'), 'Luca Conti')
    _check(r, 'Intestatario', 'ma il nome del file resta quello del cliente',
           nome_file('anna'), 'Anna')
    _check(r, 'Intestatario', 'senza intestatario: la fattura va al cliente',
           intestata_a('sara'), 'Sara Bernasconi')
    _check(r, 'Intestatario', 'campo con soli spazi vale come vuoto',
           (lambda: (con.execute("UPDATE clients SET intestatario='   ' WHERE key='sara'"),
                     intestata_a('sara'))[1])(), 'Sara Bernasconi')
    con.close()


def _test_banca(r):
    _test_intestatario(r)
    """Leggere l'estratto e accostarlo: qui un errore costa caro, si prova bene."""
    import os
    import sqlite3
    import tempfile
    from . import banca as B
    from . import db as _db
    from .db import SCHEMA

    _check(r, 'Banca', "1'234.50 svizzero", B._importo("1'234.50"), 123450)
    _check(r, 'Banca', '1.234,50 all\'italiana', B._importo('1.234,50'), 123450)
    _check(r, 'Banca', 'importo negativo riconosciuto', B._importo('-45.00'), -4500)
    _check(r, 'Banca', 'cella vuota non è zero', B._importo('  '), None)
    _check(r, 'Banca', 'data svizzera 02.08.2026', B._data('02.08.2026'), '2026-08-02')
    _check(r, 'Banca', 'data ISO 2026-08-02', B._data('2026-08-02'), '2026-08-02')

    cartella = tempfile.mkdtemp()
    with open(os.path.join(cartella, 'e.csv'), 'w', encoding='utf-8') as f:
        f.write(CSV_PROVA)
    with open(os.path.join(cartella, 'e.xml'), 'w', encoding='utf-8') as f:
        f.write(CAMT_PROVA)
    mov, problemi = B.leggi_cartella(cartella)

    _check(r, 'Banca', 'nessun problema sui due formati standard', problemi, [])
    _check(r, 'Banca', 'legge sia il CSV sia il camt', len(mov), 4)
    _check(r, 'Banca', 'gli addebiti non entrano (non sono incassi)',
           [m for m in mov if m['importo_cents'] < 0], [])
    _check(r, 'Banca', "l'addebito della cassa malati è escluso",
           any('Krankenkasse' in m['descrizione'] for m in mov), False)
    _check(r, 'Banca', 'il camt dà il nome di chi ha pagato',
           [m['nome'] for m in mov if m['importo_cents'] == 180000], ['Bruno Keller'])
    _check(r, 'Banca', 'lo stesso file letto due volte non raddoppia i movimenti',
           len(B.leggi_cartella(cartella)[0]), 4)

    _check(r, 'Banca', 'la dieresi scritta «ue» dalla banca si riconosce',
           B.somiglianza_nome('Zahlung Mueller Petra', 'Petra Müller'), 1.0)
    _check(r, 'Banca', 'la dieresi scritta senza nulla si riconosce',
           B.somiglianza_nome('Zahlung Muller Petra', 'Petra Müller'), 1.0)
    _check(r, 'Banca', 'una causale anonima non somiglia a nessuno',
           B.somiglianza_nome('Gutschrift unbekannt', 'Sofia Ferrari'), 0.0)

    con = sqlite3.connect(':memory:')
    con.row_factory = sqlite3.Row
    con.executescript(SCHEMA)
    _db._migrate(con)
    ins = ('INSERT INTO invoices(number, client_name, date, year, total_cents, status, '
           'source) VALUES(?,?,?,?,?,?,"app")')
    con.execute(ins, (82, 'Sofia Ferrari', '2026-07-24', 2026, 11000, 'emessa'))
    con.execute(ins, (83, 'Chiara De Santis', '2026-07-25', 2026, 11000, 'emessa'))
    con.execute(ins, (60, 'Nadia Rossi', '2026-07-20', 2026, 11000, 'pagata'))
    con.commit()

    versamento = next(m for m in mov if m['data'] == '2026-08-02')
    cand = B.candidati_per(con, versamento)
    _check(r, 'Banca', 'il nome nella causale mette la fattura giusta in cima',
           cand[0]['inv']['number'], 82)
    _check(r, 'Banca', 'la fattura col nome giusto è «probabile»',
           cand[0]['grado'], B.PROBABILE)
    _check(r, 'Banca', "quelle con lo stesso importo ma altro nome restano «possibile»",
           {c['grado'] for c in cand[1:]}, {B.POSSIBILE})
    # le fatture gia' spuntate a mano restano candidate: confermarle non cambia
    # lo stato ma scrive la data vera del versamento
    pagata = next(c for c in cand if c['inv']['number'] == 60)
    _check(r, 'Banca', 'una fattura già pagata a mano resta proponibile',
           pagata['aperta'], False)
    _check(r, 'Banca', 'e la riga lo dice, invece di far credere a un incasso nuovo',
           'già segnata pagata' in pagata['perche'], True)
    _check(r, 'Banca', 'a parità di indizi la fattura ancora aperta viene prima',
           cand[1]['aperta'], True)

    anonimo = next(m for m in mov if 'unbekannt' in m['descrizione'])
    _check(r, 'Banca', 'senza nome nessun candidato è «probabile»',
           [c for c in B.candidati_per(con, anonimo) if c['grado'] != B.POSSIBILE], [])
    _check(r, 'Banca', 'due candidati deboli NON fanno una proposta automatica',
           B.proposte(con, [anonimo])[0]['chiaro'], False)
    _check(r, 'Banca', 'un candidato che stacca gli altri sì',
           B.proposte(con, [versamento])[0]['chiaro'], True)

    # la data della fattura scritta nella causale (alcuni clienti fanno cosi')
    _check(r, 'Banca', 'la data citata «21-04-26» si legge',
           '2026-04-21' in B.date_citate('INVOICE 21-04-26'), True)
    _check(r, 'Banca', 'anche scritta «04.09.25»',
           '2025-09-04' in B.date_citate('INVOICE 04.09.25'), True)
    _check(r, 'Banca', 'una data impossibile non si inventa',
           B.date_citate('INVOICE 45-99-26'), set())
    con.execute(ins, (95, 'Tizia Rossi', '2026-07-20', 2026, 50000, 'emessa'))
    con.execute(ins, (96, 'Tizia Rossi', '2026-07-28', 2026, 50000, 'emessa'))
    con.commit()
    due_uguali = dict(versamento, importo_cents=50000,
                      descrizione='Accredito Tizia Rossi INVOICE 28-07-26')
    cd = B.candidati_per(con, due_uguali)
    _check(r, 'Banca', 'fra due mensilità identiche vince quella con la data citata',
           cd[0]['inv']['number'], 96)
    _check(r, 'Banca', 'e diventa una proposta da confermare con un click',
           B.proposte(con, [due_uguali])[0]['chiaro'], True)

    # una mensilita' gia' spuntata a mano ma del mese giusto deve battere quella
    # ancora aperta del mese dopo (caso Nadia: versamento del 30.06)
    con.execute(ins, (97, 'Tizia Rossi', '2026-06-20', 2026, 33000, 'pagata'))
    con.execute(ins, (98, 'Tizia Rossi', '2026-07-01', 2026, 33000, 'emessa'))
    con.commit()
    fine_mese = dict(versamento, data='2026-06-30', importo_cents=33000,
                     descrizione='Accredito Tizia Rossi')
    cf = B.candidati_per(con, fine_mese)
    _check(r, 'Banca', 'il versamento va alla fattura emessa PRIMA, non a quella dopo',
           cf[0]['inv']['number'], 97)
    _check(r, 'Banca', "e non basta che l'altra sia ancora aperta per scavalcarla",
           cf[0]['aperta'], False)
    con.execute('DELETE FROM invoices WHERE number IN (97,98)')
    con.commit()

    # un bonifico che paga due fatture insieme
    con.execute(ins, (90, 'Tizia Rossi', '2026-07-10', 2026, 200000, 'emessa'))
    con.execute(ins, (91, 'Tizia Rossi', '2026-07-10', 2026, 135000, 'emessa'))
    con.execute(ins, (92, 'Caio Bianchi', '2026-07-10', 2026, 335000, 'emessa'))
    con.commit()
    insieme = dict(versamento, importo_cents=335000,
                   descrizione='Accredito Tizia Rossi Via Roma 1')
    g = B.gruppi_per(con, insieme)
    _check(r, 'Banca', 'due fatture che insieme fanno il bonifico si trovano',
           sorted(i['number'] for i in g[0]['fatture']) if g else [], [90, 91])
    _check(r, 'Banca', 'il gruppo si cerca solo fra le fatture di chi ha pagato',
           any(92 in [i['number'] for i in x['fatture']] for x in g), False)
    _check(r, 'Banca', 'se una fattura da sola basta, non si cercano gruppi',
           B.proposte(con, [dict(versamento, importo_cents=335000,
                                 descrizione='Accredito Caio Bianchi')])[0]['gruppi'], [])

    # il numero della fattura scritto a mano: e' l'unica via per i pagamenti
    # che l'app non puo' proporre (arrivati mesi dopo, o di importo diverso)
    doppio = con.execute('SELECT COUNT(*) c FROM invoices WHERE number=82').fetchone()['c']
    _check(r, 'Banca', 'partenza pulita per la prova del numero doppio', doppio, 1)
    con.execute(ins, (82, 'Caio Bianchi', '2026-07-24', 2026, 11000, 'emessa'))
    con.commit()
    stesso = con.execute('SELECT * FROM invoices WHERE number=82').fetchall()
    _check(r, 'Banca', 'due fatture possono avere lo stesso numero (succede)',
           len(stesso), 2)
    suoi = [t for t in stesso
            if B.somiglianza_nome('Accredito Sofia Ferrari', t['client_name']) >= 0.5]
    _check(r, 'Banca', 'il nome di chi versa scioglie il numero doppio',
           [t['client_name'] for t in suoi], ['Sofia Ferrari'])
    con.execute('DELETE FROM invoices WHERE number=82 AND client_name="Caio Bianchi"')
    con.commit()

    fuori_finestra = dict(versamento, data='2027-08-02')
    _check(r, 'Banca', 'un versamento di un anno dopo non si attacca a niente',
           B.candidati_per(con, fuori_finestra), [])

    # il riferimento non si deduce dal numero: sarebbe un "certo" inventato
    con_rif = dict(versamento, riferimento='000000000000000000000000082')
    _check(r, 'Banca', 'nessun «certo» finché le fatture non hanno un riferimento vero',
           [c for c in B.candidati_per(con, con_rif) if c['grado'] == B.CERTO], [])
    con.close()

    import shutil
    shutil.rmtree(cartella, ignore_errors=True)